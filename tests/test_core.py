# -*- coding: utf-8 -*-
import os
import sys
import unittest
import tempfile
import hashlib
import base64
import datetime
import shutil
import subprocess
from pathlib import Path

from docx import Document

PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_DIR)

from tools.credit_code import generate_batch, validate_code
from tools.id_documents import (
    DOCUMENT_TYPES, generate_personal_batch, resident_id_age,
    resident_id_gender, validate_personal_document,
)
from tools.sql_tool import (
    build_sql_package, classify_sql_type, export_sql_package,
    deduplicate_sql_statements,
    generate_file_header, generate_reverse_sql, generate_verification_sql,
    split_mixed_sql, split_statements, strip_comments, validate_oracle_sql,
    validate_oracle_sql_detailed,
)
from config import DEFAULT_SETTINGS, DEFAULT_SYSTEMS, normalize_settings
from tools.docx_updater import (
    DATA_TABLE_HEADERS, deduplicate_sql, detect_existing_changes,
    filter_docx_structure_sql, parse_sql,
    process, refined_docx_path,
)
from tools.docx_template_registry import TEMPLATE_PROFILES, match_document_template
from tools.vin_generator import generate_vin_batch, validate_vin
from tools.gateway_crypto import KEYS, decrypt_gateway_payload
from tools.ops_commands import (
    COMMANDS, build_command, contains_forbidden_delete, infer_risk,
    load_custom_commands, output_guide, save_custom_commands, search_commands,
)
from tools.json_viewer import (
    format_json_text, json_path_child, node_json_text, node_value_text,
    parse_json_text, search_json_nodes,
)
from tools.daily_reports import is_reminder_due, normalize_reminder
from tools.personal_knowledge import (
    entry_fingerprint, export_word_entry, export_workbook_entry, extract_word_entry,
    organize_content, search_entries,
)
from tools.requirements import (
    classify_requirement, daily_template, requirement_from_text,
    merge_working_copies, requirement_search_text,
)
from tools.svn_workspace import (
    add_text_file, checkout, commit_working_copy, infer_online_month,
    infer_record_kind, run_svn, scan_working_copies, svn_status,
    update_working_copy,
)
from gmssl import sm2, sm4


class CreditCodeTests(unittest.TestCase):
    def test_batch_is_unique_and_valid(self):
        codes = generate_batch(200)
        self.assertEqual(len(codes), 200)
        self.assertEqual(len(set(codes)), 200)
        self.assertTrue(all(validate_code(code) for code in codes))


class PrivateWorkspaceTests(unittest.TestCase):
    def test_lazy_knowledge_organization_classifies_and_searches(self):
        text = (
            'Oracle 测试数据库\nSERVICE_NAME = simutfdb\n用户名 sitautocore\n\n'
            '====================\n\n'
            '接口联调说明\nhttps://gateway.example/api\n请求报文示例'
        )
        entries = organize_content(text, source='测试粘贴')
        self.assertEqual(len(entries), 2)
        self.assertEqual(entries[0]['category'], 'database')
        self.assertEqual(entries[1]['category'], 'interface')
        self.assertTrue(entries[0]['sensitive'])
        self.assertEqual(search_entries(entries, 'simutfdb'), [entries[0]])
        self.assertEqual(search_entries(entries, '接口', 'interface'), [entries[1]])

    def test_structured_workbook_rows_are_searchable_and_stable(self):
        entry = {
            'content_type': 'workbook_sheet', 'title': '测试表 · 内网',
            'source': '测试表.xlsx', 'sheet_name': '内网', 'category': 'server',
            'content': '', 'rows': [['编号', '主机'], ['1', 'server-a'], ['2', 'server-b']],
        }
        self.assertEqual(search_entries([entry], 'server-b'), [entry])
        self.assertEqual(search_entries([entry], 'server-b', 'database'), [])
        self.assertEqual(entry_fingerprint(entry), entry_fingerprint(dict(entry)))

    def test_visible_excel_export_and_word_type(self):
        from openpyxl import load_workbook
        entry = {
            'sheet_name': '内网', 'rows': [['编号', '主机', '备注'], ['1', 'alpha', 'A'], ['2', 'beta', 'B']],
            'column_count': 3, 'column_widths': [8, 20, 12], 'header_rows': [0], 'cell_styles': {},
        }
        with tempfile.TemporaryDirectory() as folder:
            output = os.path.join(folder, 'visible.xlsx')
            export_workbook_entry(entry, output, [0, 2], [0, 1])
            workbook = load_workbook(output, data_only=True)
            self.assertEqual(workbook.active.max_row, 2)
            self.assertEqual(workbook.active.max_column, 2)
            self.assertEqual(workbook.active['B2'].value, 'beta')
            workbook.close()
            word_path = os.path.join(folder, 'requirement.docx')
            document = Document(); document.add_paragraph('需求说明'); document.save(word_path)
            word_entry = extract_word_entry(word_path)
            self.assertEqual(word_entry['file_type'], 'WORD')
            self.assertEqual(word_entry['content_type'], 'word_document')
            exported_word = os.path.join(folder, 'edited.docx')
            word_entry['content'] = '编辑后的需求说明'
            word_entry['document_html'] = '<p>编辑后的需求说明</p>'
            export_word_entry(word_entry, exported_word)
            self.assertEqual(Document(exported_word).paragraphs[0].text, '编辑后的需求说明')

    def test_daily_reminder_and_requirement_template(self):
        settings = normalize_reminder({'enabled': True, 'time': '17:30', 'last_reminder_date': ''})
        now = datetime.datetime(2026, 7, 16, 17, 31)
        self.assertTrue(is_reminder_due(settings, now))
        requirement = requirement_from_text(
            'REQ-20260716-0001 新增投保接口，需要周边系统同步升级并执行 ALTER TABLE。',
            '需求说明.md',
        )
        self.assertEqual(requirement['category'], '接口联动')
        self.assertTrue(requirement['has_sql'])
        self.assertTrue(requirement['needs_peripheral_upgrade'])
        self.assertIn('req-20260716-0001', requirement_search_text(requirement))
        template = daily_template(requirement)
        self.assertIn('需求', template['completed'])
        self.assertIn('含 SQL', template['notes'])
        self.assertIn('周边系统', template['notes'])

    def test_reminder_settings_file_roundtrip_shared_store(self):
        """设置页与日报页共用 daily_report_settings.json。"""
        import tempfile
        from tools.daily_reports import load_reminder_settings, save_reminder_settings
        with tempfile.TemporaryDirectory() as folder:
            path = os.path.join(folder, 'daily_report_settings.json')
            saved = save_reminder_settings(
                {'enabled': True, 'time': '16:40', 'last_reminder_date': ''},
                path=path,
            )
            self.assertTrue(saved['enabled'])
            self.assertEqual(saved['time'], '16:40')
            loaded = load_reminder_settings(path)
            self.assertEqual(loaded['time'], '16:40')
            self.assertTrue(loaded['enabled'])
            # 改时间后清空 last_reminder_date 由 UI 层负责；文件层原样读写
            again = save_reminder_settings(
                {'enabled': False, 'time': '08:05', 'last_reminder_date': '2026-07-21'},
                path=path,
            )
            self.assertFalse(again['enabled'])
            self.assertEqual(again['time'], '08:05')
            self.assertEqual(load_reminder_settings(path)['last_reminder_date'], '2026-07-21')

    def test_requirement_classification_covers_common_types(self):
        self.assertEqual(classify_requirement('修复 BUG：保存时报错'), '缺陷优化')
        self.assertEqual(classify_requirement('新增字段并执行 DDL SQL'), '数据变更')


class SvnWorkspaceTests(unittest.TestCase):
    def test_month_and_bug_are_inferred_from_folder_names(self):
        self.assertEqual(infer_online_month(r'C:\需求\2026-02车险需求\REQ-001'), '2026-02')
        self.assertEqual(infer_online_month(r'C:\需求\3月上线\REQ-002', default_year=2027), '2027-03')
        self.assertEqual(infer_online_month(r'C:\2026年车险需求\六月\REQ-003'), '2026-06')
        self.assertEqual(infer_record_kind(r'2月\BUG-1234-保存异常'), 'BUG')
        self.assertEqual(infer_record_kind(r'2月\REQ-1234-新增功能'), '需求')

    def test_plain_requirement_folders_are_grouped_and_sorted_by_modified_time(self):
        with tempfile.TemporaryDirectory() as folder:
            root = os.path.join(folder, '2026年车险需求')
            older = os.path.join(root, '六月', 'REQ-OLD')
            newer = os.path.join(root, '六月', 'REQ-NEW')
            february = os.path.join(root, '2月', 'REQ-FEB')
            for path in (older, newer, february):
                os.makedirs(path)
            files = [os.path.join(older, 'old.txt'), os.path.join(newer, 'new.txt'), os.path.join(february, 'feb.txt')]
            for path in files:
                with open(path, 'w', encoding='utf-8') as stream:
                    stream.write(path)
            os.utime(files[0], (1000, 1000))
            os.utime(files[1], (3000, 3000))
            os.utime(files[2], (4000, 4000))

            copies = scan_working_copies(root)
            self.assertEqual([item['online_month'] for item in copies], ['2026-06', '2026-06', '2026-02'])
            self.assertEqual([os.path.basename(item['local_path']) for item in copies], ['REQ-NEW', 'REQ-OLD', 'REQ-FEB'])
            self.assertTrue(all(item['workspace_kind'] == 'folder' for item in copies))
            self.assertTrue(all(item['file_count'] == 1 for item in copies))

    def test_real_local_svn_checkout_add_commit_update_and_scan(self):
        with tempfile.TemporaryDirectory() as folder:
            repository = os.path.join(folder, 'repository')
            svnadmin = os.path.join(os.path.dirname(shutil.which('svn')), 'svnadmin.exe')
            subprocess.run([svnadmin, 'create', repository], check=True)
            url = Path(repository).as_uri()
            run_svn(['mkdir', f'{url}/trunk', '-m', 'init'])
            checkout_path = os.path.join(folder, '1月车险需求', 'REQ-20260717-0001')
            checkout(f'{url}/trunk', checkout_path)
            created = add_text_file(checkout_path, '文档/需求说明.md', '新增投保功能')
            self.assertTrue(os.path.isfile(created))
            self.assertFalse(svn_status(checkout_path)['clean'])
            committed = commit_working_copy(checkout_path, '新增需求说明')
            self.assertTrue(committed['svn_revision'])
            self.assertTrue(svn_status(checkout_path)['clean'])
            updated = update_working_copy(checkout_path)
            self.assertEqual(updated['svn_url'], f'{url}/trunk')
            copies = scan_working_copies(folder)
            self.assertEqual(len(copies), 1)
            self.assertEqual(copies[0]['online_month'], f'{datetime.date.today().year}-01')
            requirements, added, changed = merge_working_copies([], copies)
            self.assertEqual((added, changed), (1, 0))
            self.assertEqual(requirements[0]['svn_url'], f'{url}/trunk')


class PersonalDocumentTests(unittest.TestCase):
    def test_all_supported_personal_documents_are_unique_and_format_valid(self):
        for kind in DOCUMENT_TYPES:
            documents = generate_personal_batch(kind, 200)
            self.assertEqual(len(documents), 200, kind)
            self.assertEqual(len(set(documents)), 200, kind)
            self.assertTrue(all(validate_personal_document(kind, item) for item in documents), kind)

    def test_resident_id_rejects_bad_date_and_check_digit(self):
        self.assertFalse(validate_personal_document('resident_id', '110101202602300011'))
        number = generate_personal_batch('resident_id', 1)[0]
        replacement = '0' if number[-1] != '0' else '1'
        self.assertFalse(validate_personal_document('resident_id', number[:-1] + replacement))

    def test_personal_document_shapes(self):
        self.assertRegex(generate_personal_batch('passport', 1)[0], r'^E\d{8}$')
        self.assertRegex(generate_personal_batch('military_officer', 1)[0], r'^军字第\d{8}号$')
        self.assertRegex(generate_personal_batch('armed_police', 1)[0], r'^武字第\d{8}号$')

    def test_resident_ids_follow_selected_district_age_and_gender(self):
        documents = generate_personal_batch(
            'resident_id', 200, area_code='440304', min_age=35, max_age=35,
            gender='female',
        )
        self.assertTrue(all(number.startswith('440304') for number in documents))
        self.assertTrue(all(resident_id_age(number) == 35 for number in documents))
        self.assertTrue(all(resident_id_gender(number) == 'female' for number in documents))
        self.assertTrue(all(validate_personal_document('resident_id', number) for number in documents))


class SqlToolTests(unittest.TestCase):
    def test_semicolon_inside_string_is_not_split(self):
        sql = "INSERT INTO T(A) VALUES ('x;y'); UPDATE T SET A='z' WHERE ID=1;"
        self.assertEqual(len(split_statements(sql)), 2)

    def test_escaped_quote_and_comment_tokens_inside_string(self):
        sql = "INSERT INTO T(A,B) VALUES ('O''Brien', '--not comment'); -- comment\nCOMMIT;"
        clean = strip_comments(sql)
        self.assertIn("O''Brien", clean)
        self.assertIn('--not comment', clean)
        self.assertNotIn('-- comment', clean)

    def test_mixed_classification_and_split(self):
        sql = 'CREATE TABLE T(ID NUMBER); INSERT INTO T(ID) VALUES (1);'
        self.assertEqual(classify_sql_type(sql), 'MIXED')
        result = split_mixed_sql(sql)
        self.assertIn('CREATE TABLE', result['ddl'])
        self.assertIn('INSERT INTO', result['dml'])

    def test_multiple_sources_are_deduplicated_into_single_ddl_and_dml_streams(self):
        sql = (
            "-- file one\nCREATE TABLE T_BATCH(ID NUMBER);\n"
            "-- file two\ncreate table t_batch(id number);\n"
            "INSERT INTO T_BATCH(ID, NAME) VALUES (1, 'A');\n"
            "insert into t_batch(id, name) values (1, 'A');\n"
            "INSERT INTO T_BATCH(ID, NAME) VALUES (2, 'a');"
        )
        unique, duplicates = deduplicate_sql_statements(sql)
        self.assertEqual(len(duplicates), 2)
        self.assertEqual(len(split_statements(unique)), 3)
        artifacts = build_sql_package(unique, DEFAULT_SYSTEMS[1], '生产环境', '20260629')
        upgrades = [item for item in artifacts if item['kind'] == 'upgrade']
        self.assertEqual([item['category'] for item in upgrades], ['DDL', 'DML'])
        self.assertEqual(upgrades[0]['content'].upper().count('CREATE TABLE T_BATCH'), 1)
        self.assertIn("VALUES (1, 'A')", upgrades[1]['content'])
        self.assertIn("VALUES (2, 'a')", upgrades[1]['content'])

    def test_insert_rollback_preserves_value_case(self):
        sql = "INSERT INTO T(ID, NAME) VALUES (1, 'LiHaopeng');"
        rollback = generate_reverse_sql(sql)
        self.assertIn("NAME='LiHaopeng'", rollback)

    def test_rollback_drops_index_before_table(self):
        """升级 CREATE TABLE + CREATE INDEX 时，回滚不得先 DROP TABLE 再 DROP INDEX。"""
        sql = (
            "CREATE TABLE T_DEMO(ID NUMBER);\n"
            "CREATE INDEX IDX_T_DEMO_ID ON T_DEMO(ID);\n"
            "CREATE UNIQUE INDEX UK_T_DEMO_CODE ON T_DEMO(CODE);"
        )
        rollback = generate_reverse_sql(sql).upper()
        idx_pos = rollback.find('DROP INDEX IDX_T_DEMO_ID')
        uk_pos = rollback.find('DROP INDEX UK_T_DEMO_CODE')
        table_pos = rollback.find('DROP TABLE T_DEMO')
        self.assertGreaterEqual(idx_pos, 0)
        self.assertGreaterEqual(uk_pos, 0)
        self.assertGreaterEqual(table_pos, 0)
        self.assertLess(idx_pos, table_pos)
        self.assertLess(uk_pos, table_pos)

    def test_rollback_order_with_insert_after_ddl(self):
        sql = (
            "CREATE TABLE T_DEMO(ID NUMBER);\n"
            "CREATE INDEX IDX_T_DEMO ON T_DEMO(ID);\n"
            "INSERT INTO T_DEMO(ID) VALUES (1);"
        )
        rollback = generate_reverse_sql(sql).upper()
        del_pos = rollback.find('DELETE FROM T_DEMO')
        idx_pos = rollback.find('DROP INDEX IDX_T_DEMO')
        table_pos = rollback.find('DROP TABLE T_DEMO')
        self.assertLess(del_pos, idx_pos)
        self.assertLess(idx_pos, table_pos)

    def test_valid_ddl_has_no_false_semicolon_warning(self):
        self.assertEqual(validate_oracle_sql('CREATE TABLE T(ID NUMBER);'), [])

    def test_update_without_where_is_warned(self):
        warnings = validate_oracle_sql("UPDATE T SET NAME='X';")
        self.assertTrue(any('without WHERE' in warning for warning in warnings))

    def test_lightweight_validation_detects_common_structure_errors(self):
        sql = "INSERT T_USER(ID) VALUE (1); UPDATE T_USER NAME='X'; CREATE TABLE ;"
        issues = validate_oracle_sql_detailed(sql)
        codes = {item['code'] for item in issues}
        self.assertIn('insert_into', codes)
        self.assertIn('update_set', codes)
        self.assertIn('create_table_name', codes)

    def test_lightweight_validation_warns_about_mysql_dialect(self):
        issues = validate_oracle_sql_detailed('SELECT * FROM `T_USER` LIMIT 1;')
        self.assertTrue(any(item['code'] == 'dialect' for item in issues))

    def test_sample_header_and_delivery_paths(self):
        # 用例内用中性作者名，避免把真实姓名写进断言
        system = dict(DEFAULT_SYSTEMS[1])
        system['script_author'] = '示例作者'
        self.assertEqual(
            generate_file_header(system, '模拟环境'),
            '---- 地址：10.128.23.211\n---- sid： simutfdb\n---- 用户名：sitecif\n\n',
        )
        sql = "CREATE TABLE T_DEMO(ID NUMBER); INSERT INTO T_DEMO(ID) VALUES (1);"
        artifacts = build_sql_package(sql, system, '生产环境', '20260629')
        paths = [item['relative_path'].replace('\\', '/') for item in artifacts]
        self.assertIn('20260629/生产环境/DDL/客户信息平台-张小龙/升级SQL/示例作者-【ECIF】升级SQL.sql', paths)
        self.assertIn('20260629/生产环境/DML/客户信息平台-张小龙/回滚SQL/示例作者-【ECIF】回滚SQL.sql', paths)
        self.assertIn('20260629/验证SQL/客户信息平台-张小龙/示例作者-【ECIF】验证SQL.sql', paths)

    def test_update_rollback_marks_manual_original_values(self):
        rollback = generate_reverse_sql("UPDATE T_USER SET NAME='NEW', FLAG='1' WHERE ID=9;")
        self.assertIn('[必须人工补充]', rollback)
        self.assertIn('NAME = <升级前原值>', rollback)
        self.assertIn('WHERE ID=9', rollback)

    def test_verification_contains_prechange_backup_and_postcheck(self):
        verification = generate_verification_sql(
            "UPDATE T_USER SET NAME='NEW' WHERE ID=9; INSERT INTO T_LOG(ID, NAME) VALUES (1, 'A');"
        )
        self.assertIn('升级前原值留存', verification)
        self.assertIn('UPDATE 执行后验证', verification)
        self.assertIn("SELECT * FROM T_LOG WHERE ID=1 AND NAME='A';", verification)

    def test_delete_alias_is_preserved_in_verification(self):
        verification = generate_verification_sql("DELETE FROM PRPDCODE p WHERE p.CODETYPE='X';")
        self.assertIn("SELECT * FROM PRPDCODE p WHERE p.CODETYPE='X';", verification)

    def test_existing_connection_header_is_not_duplicated(self):
        system = DEFAULT_SYSTEMS[1]
        source = (
            generate_file_header(system, '模拟环境') + 'INSERT INTO T_LOG(ID) VALUES (1);\n'
            + generate_file_header(system, '生产环境') + 'INSERT INTO T_LOG(ID) VALUES (2);'
        )
        upgrade = next(item for item in build_sql_package(source, system, '生产环境', '20260629') if item['kind'] == 'upgrade')
        self.assertEqual(upgrade['content'].count('---- 地址：'), 1)

    def test_package_export_uses_utf8_bom(self):
        with tempfile.TemporaryDirectory() as folder:
            paths = export_sql_package(
                folder, 'INSERT INTO T_LOG(ID) VALUES (1);',
                DEFAULT_SYSTEMS[1], '模拟环境', '20260629',
            )
            self.assertEqual(len(paths), 3)
            for path in paths:
                with open(path, 'rb') as stream:
                    self.assertEqual(stream.read(3), b'\xef\xbb\xbf')


class VinTests(unittest.TestCase):
    def test_china_vins_are_unique_and_valid(self):
        vins = generate_vin_batch(200, 2026)
        self.assertEqual(len(vins), 200)
        self.assertTrue(all(vin.startswith('L') and validate_vin(vin) for vin in vins))


class GatewayCryptoTests(unittest.TestCase):
    @staticmethod
    def _vector(environment, direction):
        key = '1234567890abcdef'
        plain = '{"code":200,"message":"PengTools"}'
        cipher = sm4.CryptSM4()
        cipher.set_key(key.encode('utf-8'), sm4.SM4_ENCRYPT)
        payload = base64.b64encode(cipher.crypt_cbc(key.encode('utf-8'), plain.encode('utf-8'))).decode('ascii')
        pair = KEYS[environment][direction]
        crypt = sm2.CryptSM2(private_key=pair['private'], public_key=pair['public'], mode=1)
        encrypted_key = crypt.encrypt(key.encode('utf-8')).hex()
        if direction == 'response':
            encrypted_key = '04' + encrypted_key
        return encrypted_key, payload, plain

    def test_request_and_response_match_gateway_protocol(self):
        for environment in (1, 3):
            for direction in ('request', 'response'):
                key, payload, expected = self._vector(environment, direction)
                self.assertEqual(decrypt_gateway_payload(direction, environment, key, payload), expected)

    def test_invalid_environment_is_rejected(self):
        with self.assertRaisesRegex(ValueError, '环境'):
            decrypt_gateway_payload('request', 9, '00', 'AA==')


class JsonViewerLogicTests(unittest.TestCase):
    def test_format_keeps_chinese_boolean_and_null(self):
        formatted = format_json_text('{"姓名":"示例用户","ok":true,"value":null}')
        self.assertIn('"姓名": "示例用户"', formatted)
        self.assertIn('"ok": true', formatted)
        self.assertIn('"value": null', formatted)

    def test_json_path_supports_array_and_special_keys(self):
        self.assertEqual(json_path_child('$.data', 0), '$.data[0]')
        self.assertEqual(json_path_child('$', 'user-name'), "$['user-name']")
        self.assertEqual(json_path_child('$', "a'b"), "$['a\\'b']")

    def test_search_matches_key_path_and_value(self):
        data = parse_json_text('{"data":{"users":[{"name":"demo_user"}]}}')
        self.assertEqual(search_json_nodes(data, 'users')[0][0], '$.data.users')
        self.assertEqual(search_json_nodes(data, 'demo_user')[0][0], '$.data.users[0].name')
        self.assertTrue(search_json_nodes(data, '$.data.users[0]'))

    def test_node_copy_text_distinguishes_value_and_json(self):
        self.assertEqual(node_value_text('hello'), 'hello')
        self.assertEqual(node_value_text(True), 'true')
        self.assertIn('\n', node_json_text({'a': 1}))

    def test_invalid_json_reports_line_and_column(self):
        with self.assertRaisesRegex(ValueError, '第 2 行，第 1 列'):
            parse_json_text('{"a": 1,\n}')


class OperationsCommandTests(unittest.TestCase):
    def test_builtin_library_has_no_file_delete_commands(self):
        self.assertGreaterEqual(len(COMMANDS), 70)
        for command in COMMANDS:
            self.assertFalse(contains_forbidden_delete(command['command']), command['command'])
            self.assertFalse(contains_forbidden_delete(command['template']), command['template'])

    def test_fuzzy_search_ps_ef_and_chinese(self):
        self.assertEqual(search_commands('ps -ef', limit=1)[0]['command'], 'ps -ef')
        self.assertEqual(search_commands('日志截取', limit=1)[0]['workflow'], 'log_extract')

    def test_log_workflow_guides_find_then_extract(self):
        command = next(item for item in COMMANDS if item.get('workflow') == 'log_extract')
        first = build_command(command, {
            'search_root': '/var/log', 'file_pattern': '*.log', 'days': '7',
            'log_path': '', 'keyword': 'ERROR', 'context': '20',
            'output_file': '/tmp/result.log',
        })
        self.assertIn('第 1 步', first)
        self.assertNotIn('grep -n', first)
        second = build_command(command, {
            'search_root': '/var/log', 'file_pattern': '*.log', 'days': '7',
            'log_path': '/var/log/app/app.log', 'keyword': 'order failed',
            'context': '20', 'output_file': '/tmp/result.log',
        })
        self.assertIn('第 2 步', second)
        self.assertIn("'order failed'", second)
        self.assertIn('> /tmp/result.log', second)

    def test_state_changing_custom_command_is_forced_dangerous(self):
        self.assertEqual(infer_risk('systemctl restart nginx'), 'danger')
        self.assertEqual(infer_risk('ps -ef'), 'safe')

    def test_output_guide_explains_ps_columns(self):
        command = search_commands('ps -ef', limit=1)[0]
        guide = output_guide(command, 'zh')
        self.assertIn('PID', guide)
        self.assertIn('PPID', guide)
        self.assertIn('CMD', guide)

    def test_only_custom_commands_are_persisted_and_delete_commands_rejected(self):
        with tempfile.TemporaryDirectory() as folder:
            path = os.path.join(folder, 'custom.json')
            custom = dict(COMMANDS[0], builtin=False, command='echo health', template='echo health')
            save_custom_commands([dict(COMMANDS[1], builtin=True), custom], path)
            loaded = load_custom_commands(path)
            self.assertEqual(len(loaded), 1)
            self.assertFalse(loaded[0]['builtin'])
            bad = dict(custom, command='rm -rf /tmp/demo', template='rm -rf /tmp/demo')
            with self.assertRaisesRegex(ValueError, '删除'):
                save_custom_commands([bad], path)


class SettingsTests(unittest.TestCase):
    def test_settings_are_normalized_to_safe_ranges(self):
        settings = normalize_settings({
            'font_size': 99, 'floating_opacity': 1,
            'copy_feedback_ms': 99999, 'default_language': 'xx',
            'close_ask_each_time': 0, 'close_default_action': 'unknown',
            'keep_awake_enabled': 1, 'keep_awake_interval_minutes': 999,
        })
        self.assertEqual(settings['font_size'], 18)
        self.assertEqual(settings['floating_opacity'], 45)
        self.assertEqual(settings['copy_feedback_ms'], 5000)
        self.assertEqual(settings['default_language'], 'zh')
        self.assertFalse(settings['close_ask_each_time'])
        self.assertEqual(settings['close_default_action'], 'minimize')
        self.assertTrue(settings['keep_awake_enabled'])
        self.assertEqual(settings['keep_awake_interval_minutes'], 60)
        self.assertTrue(DEFAULT_SETTINGS['floating_always_on_top'])
        self.assertEqual(DEFAULT_SETTINGS['floating_shortcuts'], [10, 2, 9, 5])
        # 安测默认收紧
        self.assertTrue(DEFAULT_SETTINGS.get('security_ssl_verify', False))
        self.assertTrue(DEFAULT_SETTINGS.get('security_confirm_remote_request', False))

    def test_floating_shortcuts_are_normalized(self):
        from ui.navigation_model import normalize_floating_shortcuts
        # 去重、非法、最多 6、至少 1
        self.assertEqual(
            normalize_floating_shortcuts([10, 10, 2, 99, 9, 5, 1, 4, 6, 3]),
            [10, 2, 9, 5, 1, 4],
        )
        self.assertEqual(normalize_floating_shortcuts([]), [10, 2, 9, 5])
        self.assertEqual(
            normalize_floating_shortcuts([8, 10], private_unlocked=False),
            [10],
        )
        settings = normalize_settings({'floating_shortcuts': [10, 'x', 2]})
        self.assertEqual(settings['floating_shortcuts'], [10, 2])


class DocxUpdaterTests(unittest.TestCase):
    @staticmethod
    def _make_source(path):
        document = Document()
        document.add_heading('T_CUSTOMER', level=2)
        document.add_paragraph('描述：客户表')
        table = document.add_table(rows=2, cols=7)
        for index, header in enumerate(DATA_TABLE_HEADERS):
            table.rows[0].cells[index].text = header
        for index, value in enumerate(('1', 'ID', '主键', 'NUMBER', '18', '', '')):
            table.rows[1].cells[index].text = value
        document.save(path)

    def test_parser_keeps_decimal_type_length(self):
        parsed = parse_sql('CREATE TABLE T_PRICE (AMOUNT NUMBER(12,2));')
        self.assertEqual(parsed['new_tables'][0]['columns'][0]['length'], '12,2')

    def test_parser_accepts_schema_and_single_column_syntax(self):
        parsed = parse_sql("ALTER TABLE APP.T_CUSTOMER ADD COLUMN NICK_NAME VARCHAR2(40); COMMENT ON COLUMN APP.T_CUSTOMER.NICK_NAME IS '昵称';")
        self.assertEqual(parsed['alter_adds'][0]['table_name'], 'T_CUSTOMER')
        self.assertEqual(parsed['alter_adds'][0]['columns'][0]['comment'], '昵称')

    def test_adds_field_and_new_table_to_existing_document(self):
        with tempfile.TemporaryDirectory() as folder:
            source = os.path.join(folder, 'interface.docx')
            sql_path = os.path.join(folder, 'change.sql')
            output = os.path.join(folder, 'interface_updated.docx')
            document = Document()
            document.add_heading('T_CUSTOMER', level=2)
            document.add_paragraph('描述：客户表')
            table = document.add_table(rows=2, cols=7)
            for index, header in enumerate(DATA_TABLE_HEADERS):
                table.rows[0].cells[index].text = header
            for index, value in enumerate(('1', 'ID', '主键', 'NUMBER', '18', '', '')):
                table.rows[1].cells[index].text = value
            document.save(source)
            with open(sql_path, 'w', encoding='utf-8') as stream:
                stream.write(
                    "ALTER TABLE T_CUSTOMER ADD (AGE NUMBER(3));\n"
                    "COMMENT ON COLUMN T_CUSTOMER.AGE IS '年龄';\n"
                    "CREATE TABLE T_ADDRESS (ID NUMBER(18) NOT NULL, AMOUNT NUMBER(12,2));\n"
                    "COMMENT ON TABLE T_ADDRESS IS '地址表';\n"
                )
            process(source, sql_path, output_path=output, backup=False)
            updated = Document(output)
            all_cells = '\n'.join(cell.text for table in updated.tables for row in table.rows for cell in row.cells)
            headings = [paragraph.text for paragraph in updated.paragraphs]
            self.assertIn('AGE', all_cells)
            self.assertIn('年龄', all_cells)
            self.assertIn('T_ADDRESS', headings)
            self.assertIn('12,2', all_cells)

    def test_duplicate_sql_is_reported_and_removed(self):
        sql = (
            '-- 来源文件: a.sql\nALTER TABLE T_CUSTOMER ADD (AGE NUMBER(3));\n'
            '-- 来源文件: b.sql\n  alter   table T_CUSTOMER add (AGE NUMBER(3));'
        )
        unique, duplicates = deduplicate_sql(sql)
        self.assertEqual(len(duplicates), 1)
        self.assertEqual(unique.upper().count('ALTER TABLE'), 1)

    def test_docx_filter_keeps_structure_ddl_and_rejects_dml(self):
        sql = (
            'ALTER TABLE T_CUSTOMER ADD (AGE NUMBER(3));\n'
            "INSERT INTO T_CUSTOMER(ID) VALUES (1);\n"
            "COMMENT ON COLUMN T_CUSTOMER.AGE IS '年龄';"
        )
        accepted, rejected = filter_docx_structure_sql(sql)
        self.assertIn('ALTER TABLE', accepted)
        self.assertIn('COMMENT ON COLUMN', accepted)
        self.assertEqual(len(rejected), 1)
        self.assertIn('INSERT INTO', rejected[0])

    def test_existing_field_preflight_and_safe_skip(self):
        with tempfile.TemporaryDirectory() as folder:
            source = os.path.join(folder, 'interface.docx')
            self._make_source(source)
            sql = 'ALTER TABLE T_CUSTOMER ADD (ID NUMBER(18));'
            existing = detect_existing_changes(source, sql)
            self.assertEqual(existing[0]['field'], 'ID')

            sql_path = os.path.join(folder, 'change.sql')
            output = refined_docx_path(source)
            with open(sql_path, 'w', encoding='utf-8') as stream:
                stream.write(sql + '\n' + sql)
            with open(source, 'rb') as stream:
                before = hashlib.sha256(stream.read()).hexdigest()
            report = process(source, sql_path, output_path=output, backup=False)
            with open(source, 'rb') as stream:
                after = hashlib.sha256(stream.read()).hexdigest()
            self.assertEqual(before, after)
            self.assertTrue(os.path.isfile(output))
            self.assertEqual(len(report['changes']), 0)
            updated = Document(output)
            fields = [row.cells[1].text for row in updated.tables[0].rows[1:]]
            self.assertEqual(fields.count('ID'), 1)

    def test_second_update_reuses_single_organized_document(self):
        with tempfile.TemporaryDirectory() as folder:
            source = os.path.join(folder, 'interface.docx')
            self._make_source(source)
            output = refined_docx_path(source)
            sql_path = os.path.join(folder, 'change.sql')
            with open(sql_path, 'w', encoding='utf-8') as stream:
                stream.write('ALTER TABLE T_CUSTOMER ADD (AGE NUMBER(3));')
            process(source, sql_path, output_path=output, backup=False)
            with open(sql_path, 'w', encoding='utf-8') as stream:
                stream.write('ALTER TABLE T_CUSTOMER ADD (NAME VARCHAR2(40));')
            process(output, sql_path, output_path=output, backup=False)

            docx_files = [name for name in os.listdir(folder) if name.endswith('.docx')]
            self.assertEqual(sorted(docx_files), sorted(['interface.docx', 'interface_整理后接口文档.docx']))
            updated = Document(output)
            cells = '\n'.join(cell.text for table in updated.tables for row in table.rows for cell in row.cells)
            self.assertIn('AGE', cells)
            self.assertIn('NAME', cells)

    def test_selected_generation_date_is_written(self):
        with tempfile.TemporaryDirectory() as folder:
            source = os.path.join(folder, 'interface.docx')
            output = refined_docx_path(source)
            sql_path = os.path.join(folder, 'change.sql')
            self._make_source(source)
            with open(sql_path, 'w', encoding='utf-8') as stream:
                stream.write('ALTER TABLE T_CUSTOMER ADD (AGE NUMBER(3));')
            report = process(
                source, sql_path, output_path=output, backup=False,
                update_date='2030-05-20',
            )
            self.assertEqual(report['version'], 'auto-2030-05-20')
            text = '\n'.join(paragraph.text for paragraph in Document(output).paragraphs)
            self.assertIn('2030-05-20', text)

    def test_uploaded_name_matches_latest_known_system_template(self):
        self.assertEqual(len(TEMPLATE_PROFILES), 27)
        cases = {
            '接报案数据库表结构文档V1.0_整理后接口文档.docx': ('接报案', 'V2.0.docx'),
            '理赔数据库表结构文档V1.1.docx': ('理赔', 'V2.0.docx'),
            '电网理赔数据库表结构说明文档V1.0.docx': ('电网', 'V1.2.docx'),
            '客户信息平台（ECIF）-数据库表结构说明文档.docx': ('ECIF', 'ECIF数据库表结构说明文档.docx'),
            '投诉管理数据结构文档 V1.0.1.docx': ('投诉管理系统', 'V1.0.1.docx'),
            '非车理赔中心 - 理赔数据库表结构文档V1.0.0.docx': ('非车理赔中心', 'V1.0.0.docx'),
            '6.4英大财险-业务应用-2024年车险理赔服务共享中心建设-数据设计说明书-V45.0.docx': ('车险理赔系统', 'V45.0.docx'),
        }
        for filename, (system, latest_suffix) in cases.items():
            profile = match_document_template(filename)
            self.assertIsNotNone(profile, filename)
            self.assertEqual(profile['system'], system)
            self.assertTrue(profile['template'].endswith(latest_suffix))

    def test_ocr_five_column_normal_heading_template_is_updated_safely(self):
        with tempfile.TemporaryDirectory() as folder:
            source = os.path.join(folder, 'OCR识别系统数据库表结构说明文档.docx')
            output = refined_docx_path(source)
            sql_path = os.path.join(folder, 'change.sql')
            document = Document()
            revision = document.add_table(rows=3, cols=6)
            for row in revision.rows:
                row.cells[1].merge(row.cells[2])
            for index, value in {0: '版 本', 1: '类 型', 3: '日 期', 4: '作 者', 5: '说 明'}.items():
                revision.rows[0].cells[index].text = value
            for index, value in {0: 'V1.0.0', 1: 'C', 3: '2023-12-21', 4: '原作者', 5: '创建文档'}.items():
                revision.rows[1].cells[index].text = value
            document.add_paragraph('OCRIPCONFIG')
            document.add_paragraph('描述：配置表')
            table = document.add_table(rows=2, cols=5)
            headers = ('序号', '字段名称', '字段描述', '类型', '允许空')
            for index, value in enumerate(headers):
                table.rows[0].cells[index].text = value
            for index, value in enumerate(('1', 'ID', '主键', 'NUMBER', '')):
                table.rows[1].cells[index].text = value
            document.save(source)
            with open(sql_path, 'w', encoding='utf-8') as stream:
                stream.write(
                    "ALTER TABLE OCRIPCONFIG ADD (TOKEN VARCHAR2(30));\n"
                    "COMMENT ON COLUMN OCRIPCONFIG.TOKEN IS '访问令牌';\n"
                    "CREATE TABLE OCRPNEW (ID NUMBER(18) NOT NULL, NAME VARCHAR2(50));\n"
                    "COMMENT ON TABLE OCRPNEW IS '新识别表';"
                )
            report = process(
                source, sql_path, author='Lihp', output_path=output,
                backup=False, update_date='2030-05-20',
            )
            self.assertEqual(report['template']['system'], 'OCR识别系统')
            updated = Document(output)
            revision_row = updated.tables[0].rows[2]
            self.assertEqual(
                [cell.text for cell in revision_row.cells],
                ['V1.0.1', 'A', 'A', '2030-05-20', 'Lihp',
                 '新增ocrpnew表（新识别表）\nOCRIPCONFIG表增加TOKEN字段'],
            )
            self.assertEqual(len(updated.tables[0].rows), 3)
            field_tables = [table for table in updated.tables if table.rows[0].cells[0].text == '序号']
            self.assertEqual(len(field_tables), 2)
            self.assertTrue(all(len(row.cells) == 5 for table in field_tables for row in table.rows))
            cells = '\n'.join(cell.text for table in field_tables for row in table.rows for cell in row.cells)
            self.assertIn('TOKEN', cells)
            self.assertIn('访问令牌', cells)
            self.assertIn('OCRPNEW', '\n'.join(paragraph.text for paragraph in updated.paragraphs))

    def test_two_segment_version_history_keeps_its_shape(self):
        with tempfile.TemporaryDirectory() as folder:
            source = os.path.join(folder, 'ECIF数据库表结构说明文档.docx')
            output = refined_docx_path(source)
            sql_path = os.path.join(folder, 'change.sql')
            document = Document()
            version = document.add_table(rows=2, cols=4)
            for index, value in enumerate(('版本号', '日期', '作 者', '说 明')):
                version.rows[0].cells[index].text = value
            for index, value in enumerate(('V1.4', '2026-01-01', 'Lihp', '旧版本')):
                version.rows[1].cells[index].text = value
            document.add_heading('T_CUSTOMER', level=2)
            document.add_paragraph('描述：客户表')
            table = document.add_table(rows=2, cols=7)
            for index, value in enumerate(DATA_TABLE_HEADERS):
                table.rows[0].cells[index].text = value
            for index, value in enumerate(('1', 'ID', '主键', 'NUMBER', '18', '', '')):
                table.rows[1].cells[index].text = value
            document.save(source)
            with open(sql_path, 'w', encoding='utf-8') as stream:
                stream.write('ALTER TABLE T_CUSTOMER ADD (AGE NUMBER(3));')
            report = process(source, sql_path, output_path=output, backup=False)
            self.assertEqual(report['version'], 'V1.5')

    def test_claim_revision_history_writes_merged_date_author_and_description_columns(self):
        with tempfile.TemporaryDirectory() as folder:
            source = os.path.join(folder, '理赔数据库表结构文档V1.3.docx')
            output = refined_docx_path(source)
            sql_path = os.path.join(folder, 'change.sql')
            document = Document()
            revision = document.add_table(rows=2, cols=5)
            for row in revision.rows:
                row.cells[1].merge(row.cells[2])
            for index, value in {0: '版本号', 1: '日期', 3: '作 者', 4: '说 明'}.items():
                revision.rows[0].cells[index].text = value
            for index, value in {0: 'V1.3', 1: '2023-12-14', 3: '原作者', 4: '旧版本'}.items():
                revision.rows[1].cells[index].text = value
            document.add_heading('T_CLAIM', level=2)
            document.add_paragraph('描述：理赔表')
            table = document.add_table(rows=2, cols=7)
            for index, value in enumerate(DATA_TABLE_HEADERS):
                table.rows[0].cells[index].text = value
            for index, value in enumerate(('1', 'ID', '主键', 'NUMBER', '18', '', '')):
                table.rows[1].cells[index].text = value
            document.save(source)
            with open(sql_path, 'w', encoding='utf-8') as stream:
                stream.write('ALTER TABLE T_CLAIM MODIFY (ID NUMBER(20));')
            report = process(
                source, sql_path, author='Lihp', output_path=output,
                backup=False, update_date='2030-05-20',
            )
            self.assertEqual(report['revision_layout'], 'date_merged5')
            row = Document(output).tables[0].rows[-1]
            self.assertEqual(row.cells[0].text, 'V1.4')
            self.assertEqual(row.cells[1].text, '2030-05-20')
            self.assertEqual(row.cells[2].text, '2030-05-20')
            self.assertEqual(row.cells[3].text, 'Lihp')
            self.assertIn('T_CLAIM表修改ID字段类型为NUMBER(20)', row.cells[4].text)
            self.assertFalse(row.cells[4].text.endswith('\n'))

    def test_compact_eight_column_template_uses_bracket_heading_and_non_null_column(self):
        with tempfile.TemporaryDirectory() as folder:
            source = os.path.join(folder, '投诉管理数据结构文档 V1.0.1.docx')
            output = refined_docx_path(source)
            sql_path = os.path.join(folder, 'change.sql')
            document = Document()
            revision = document.add_table(rows=2, cols=5)
            for row in revision.rows:
                row.cells[1].merge(row.cells[2])
            for index, value in {0: '版本号', 1: '日期', 3: '作 者', 4: '说 明'}.items():
                revision.rows[0].cells[index].text = value
            for index, value in {0: 'V1.0.1', 1: '2026/05/27', 3: '原作者', 4: '旧版本'}.items():
                revision.rows[1].cells[index].text = value
            document.add_heading('CMP_CASE_LOG [投诉系统日志表]', level=3)
            table = document.add_table(rows=2, cols=8)
            headers = ('#', '字段', '名称', '数据类型', '主键', '非空', '默认值', '备注说明')
            for index, value in enumerate(headers):
                table.rows[0].cells[index].text = value
            for index, value in enumerate(('1', 'ID', '主键', 'NUMBER(20)', '√', '√', '', '')):
                table.rows[1].cells[index].text = value
            document.save(source)
            with open(sql_path, 'w', encoding='utf-8') as stream:
                stream.write(
                    "ALTER TABLE CMP_CASE_LOG ADD (TRACE_ID VARCHAR2(40) NOT NULL);\n"
                    "COMMENT ON COLUMN CMP_CASE_LOG.TRACE_ID IS '追踪号';\n"
                    "CREATE TABLE CMP_NEW (ID NUMBER(18) PRIMARY KEY, NAME VARCHAR2(30));\n"
                    "COMMENT ON TABLE CMP_NEW IS '新增投诉表';"
                )
            report = process(source, sql_path, author='Lihp', output_path=output,
                             backup=False, update_date='2030-05-20')
            self.assertEqual(report['template']['system'], '投诉管理系统')
            updated = Document(output)
            self.assertIn('CMP_NEW [新增投诉表]', '\n'.join(p.text for p in updated.paragraphs))
            existing = updated.tables[1]
            added = existing.rows[-1]
            self.assertEqual(added.cells[1].text, 'TRACE_ID')
            self.assertEqual(added.cells[5].text, '√')
            new_table = updated.tables[-1]
            self.assertEqual(new_table.rows[1].cells[4].text, '√')
            self.assertEqual(updated.tables[0].rows[-1].cells[0].text, 'V1.0.2')

    def test_embedded_title_template_and_date_first_revision_are_updated(self):
        with tempfile.TemporaryDirectory() as folder:
            source = os.path.join(folder, '2024年车险理赔服务共享中心建设-数据设计说明书-V45.0.docx')
            output = refined_docx_path(source)
            sql_path = os.path.join(folder, 'change.sql')
            document = Document()
            revision = document.add_table(rows=2, cols=6)
            for index, value in enumerate(('日期', '版本', '说 明', '作者/修改人', '审核', '批准')):
                revision.rows[0].cells[index].text = value
            for index, value in enumerate(('2026/05/28', 'V45', '旧版本', '原作者', '', '')):
                revision.rows[1].cells[index].text = value
            table = document.add_table(rows=3, cols=5)
            table.rows[0].cells[0].merge(table.rows[0].cells[4])
            table.rows[0].cells[0].text = '报案表PRPLREGIST'
            for index, value in enumerate(('字段名', '字段说明', '数据类型', '长度', '能否为空')):
                table.rows[1].cells[index].text = value
            for index, value in enumerate(('ID', '主键', 'NUMBER', '18', 'false')):
                table.rows[2].cells[index].text = value
            document.save(source)
            with open(sql_path, 'w', encoding='utf-8') as stream:
                stream.write(
                    "ALTER TABLE PRPLREGIST ADD (TRACE_ID VARCHAR2(40) NOT NULL);\n"
                    "COMMENT ON COLUMN PRPLREGIST.TRACE_ID IS '追踪号';\n"
                    "CREATE TABLE PRPLNEW (ID NUMBER(18), NAME VARCHAR2(30));\n"
                    "COMMENT ON TABLE PRPLNEW IS '新增车险表';"
                )
            report = process(source, sql_path, author='Lihp', output_path=output,
                             backup=False, update_date='2030-05-20')
            self.assertEqual(report['template']['system'], '车险理赔系统')
            self.assertEqual(report['revision_layout'], 'date_first6')
            updated = Document(output)
            revision_row = updated.tables[0].rows[-1]
            self.assertEqual(revision_row.cells[0].text, '2030-05-20')
            self.assertEqual(revision_row.cells[1].text, 'V46')
            self.assertEqual(revision_row.cells[3].text, 'Lihp')
            self.assertEqual(revision_row.cells[4].text, '')
            self.assertEqual(updated.tables[-1].rows[0].cells[0].text, '新增车险表 PRPLNEW')
            self.assertEqual(updated.tables[-1].rows[1].cells[0].text, '字段名')
            self.assertEqual(updated.tables[-1].rows[2].cells[0].text, 'ID')


if __name__ == '__main__':
    unittest.main()
