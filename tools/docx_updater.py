#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
数据库表结构文档自动更新工具 v3.0 - 通用版
自动检测文档格式，兼容任意系统的 DOCX 结构文档。

支持 SQL 操作：
  CREATE TABLE        → 新增表章节
  ALTER TABLE ADD     → 在已有表中追加字段行
  ALTER TABLE MODIFY  → 修改已有字段的类型/长度
  COMMENT ON COLUMN   → 更新字段描述
  COMMENT ON TABLE    → 更新表描述

自动检测：
  - 表名所在的标题级别（Heading 1/2/3...）
  - 版本历史表的位置和格式
  - TOC 目录的样式
  - 描述段落的格式
  - 找不到对应结构则优雅跳过
"""

import sys, os, re, shutil, argparse
from datetime import date
from copy import deepcopy
from collections import Counter

from tools.docx_template_registry import match_document_template

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误：请先安装 python-docx 库：pip install python-docx")
    sys.exit(1)


# ===================== 编码检测 =====================
def detect_encoding(filepath):
    for enc in ["utf-8-sig", "utf-8", "gbk", "gb18030", "gb2312", "latin-1"]:
        try:
            with open(filepath, "r", encoding=enc) as f:
                f.read()
            return enc
        except (UnicodeDecodeError, UnicodeError):
            continue
    return "utf-8"


# ===================== SQL 解析（同 v2） =====================
SQL_NAME = r"(?:[A-Za-z_][\w$#]*\.)?[A-Za-z_][\w$#]*"


def _table_name(identifier):
    """接口文档通常使用裸表名，因此自动移除 SQL 中的 schema 前缀。"""
    return identifier.rsplit(".", 1)[-1].upper()


def parse_column_defs(body_text):
    columns = []
    lines = _split_sql_list(body_text)
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if re.match(r"^\s*(CONSTRAINT|PRIMARY|FOREIGN|UNIQUE|CHECK)\s", line, re.I):
            continue
        parts = line.split()
        if len(parts) < 2:
            continue
        col_name = parts[0].strip().upper()
        type_part = parts[1].strip()
        type_match = re.match(r"(\w+)\(?([^)]*)\)?", type_part)
        col_type = type_match.group(1).upper() if type_match else type_part.upper()
        col_length = type_match.group(2) if type_match and type_match.group(2) else ""
        rest = " ".join(parts[2:]).upper() if len(parts) > 2 else ""
        is_primary_key = "PRIMARY KEY" in rest
        is_nullable = "NOT NULL" not in rest and not is_primary_key
        default_val = ""
        dm = re.search(r"DEFAULT\s+'?([^'\s]+)'?", rest)
        if dm:
            default_val = dm.group(1)
        columns.append({
            "name": col_name, "type": col_type, "length": col_length,
            "nullable": is_nullable, "primary_key": is_primary_key,
            "default": default_val,
        })
    return columns


def _split_sql_list(text):
    """按最外层逗号拆分，保留 NUMBER(12,2) 等类型参数。"""
    parts, start, depth, in_string = [], 0, 0, False
    index = 0
    while index < len(text):
        char = text[index]
        if char == "'" and in_string and index + 1 < len(text) and text[index + 1] == "'":
            index += 2
            continue
        if char == "'":
            in_string = not in_string
        elif not in_string:
            if char == "(":
                depth += 1
            elif char == ")":
                depth = max(0, depth - 1)
            elif char == "," and depth == 0:
                parts.append(text[start:index])
                start = index + 1
        index += 1
    parts.append(text[start:])
    return parts


def split_sql_statements(sql_content):
    """Split SQL at top-level semicolons while preserving quoted text."""
    statements, start, in_string = [], 0, False
    index = 0
    while index < len(sql_content):
        char = sql_content[index]
        if char == "'" and in_string and index + 1 < len(sql_content) and sql_content[index + 1] == "'":
            index += 2
            continue
        if char == "'":
            in_string = not in_string
        elif char == ";" and not in_string:
            statement = sql_content[start:index].strip()
            if statement:
                statements.append(statement + ";")
            start = index + 1
        index += 1
    tail = sql_content[start:].strip()
    if tail:
        statements.append(tail if tail.endswith(";") else tail + ";")
    return statements


def normalize_sql_statement(statement):
    """Canonical representation used only for duplicate comparison."""
    statement = re.sub(r"(?m)^\s*--[^\n]*(?:\n|$)", "", statement)
    return re.sub(r"\s+", " ", statement.strip().rstrip(";")).casefold()


def deduplicate_sql(sql_content):
    """Return unique SQL text and duplicate statements in source order."""
    seen, unique, duplicates = set(), [], []
    for statement in split_sql_statements(sql_content):
        key = normalize_sql_statement(statement)
        if key in seen:
            duplicates.append(statement)
        else:
            seen.add(key)
            unique.append(statement)
    return "\n".join(unique), duplicates


def filter_docx_structure_sql(sql_content):
    """只保留接口结构文档能够消费的 DDL，返回 (保留文本, 被过滤语句)。"""
    accepted, rejected = [], []
    patterns = (
        r"^CREATE\s+TABLE\b",
        r"^ALTER\s+TABLE\b.*\b(ADD|MODIFY)\b",
        r"^COMMENT\s+ON\s+(TABLE|COLUMN)\b",
    )
    for statement in split_sql_statements(sql_content):
        normalized = re.sub(r"(?m)^\s*--[^\n]*(?:\n|$)", "", statement).strip()
        target = accepted if any(re.search(pattern, normalized, re.I | re.S) for pattern in patterns) else rejected
        target.append(statement)
    return "\n".join(accepted), rejected


def refined_docx_path(docx_path):
    """Deterministic output: original and one organized document side-by-side."""
    root, ext = os.path.splitext(os.path.abspath(docx_path))
    suffix = "_整理后接口文档"
    return root + ext if root.endswith(suffix) else root + suffix + (ext or ".docx")


def parse_sql(sql_content):
    result = {
        "new_tables": [], "alter_adds": [], "alter_modifies": [],
        "col_comments": [], "table_comments": [],
    }
    # CREATE TABLE
    for m in re.finditer(rf"CREATE\s+TABLE\s+({SQL_NAME})\s*\((.*?)\);", sql_content, re.DOTALL | re.I):
        table_name = _table_name(m.group(1))
        body = m.group(2)
        table_comment = ""
        qualified = rf"(?:[A-Za-z_][\w$#]*\.)?{re.escape(table_name)}"
        tm = re.search(rf"COMMENT\s+ON\s+TABLE\s+{qualified}\s+IS\s+'([^']*)'", sql_content, re.I)
        if tm:
            table_comment = tm.group(1)
        col_comments = {}
        for cm in re.finditer(rf"COMMENT\s+ON\s+COLUMN\s+{qualified}\.(\w+)\s+IS\s+'([^']*)'", sql_content, re.I):
            col_comments[cm.group(1).lower()] = cm.group(2)
        columns = parse_column_defs(body)
        for col in columns:
            col["comment"] = col_comments.get(col["name"].lower(), "")
        result["new_tables"].append({
            "table_name": table_name, "table_comment": table_comment, "columns": columns,
        })
    # ALTER TABLE ADD
    for m in re.finditer(rf"ALTER\s+TABLE\s+({SQL_NAME})\s+ADD(?:\s+COLUMN)?\s*(?:\((.*?)\)|([^;]+));", sql_content, re.DOTALL | re.I):
        table_name = _table_name(m.group(1))
        body = m.group(2) or m.group(3)
        columns = parse_column_defs(body)
        for col in columns:
            qualified = rf"(?:[A-Za-z_][\w$#]*\.)?{re.escape(table_name)}"
            cm = re.search(rf"COMMENT\s+ON\s+COLUMN\s+{qualified}\.{col['name']}\s+IS\s+'([^']*)'", sql_content, re.I)
            col["comment"] = cm.group(1) if cm else ""
        result["alter_adds"].append({"table_name": table_name, "columns": columns})
    # ALTER TABLE MODIFY
    for m in re.finditer(rf"ALTER\s+TABLE\s+({SQL_NAME})\s+MODIFY(?:\s+COLUMN)?\s*(?:\((.*?)\)|([^;]+));", sql_content, re.DOTALL | re.I):
        table_name = _table_name(m.group(1))
        body = m.group(2) or m.group(3)
        columns = parse_column_defs(body)
        result["alter_modifies"].append({"table_name": table_name, "columns": columns})
    # COMMENT ON COLUMN (独立)
    new_table_names = {t["table_name"].upper() for t in result["new_tables"]}
    for m in re.finditer(rf"COMMENT\s+ON\s+COLUMN\s+({SQL_NAME})\.(\w+)\s+IS\s+'([^']*)'", sql_content, re.I):
        table_name = _table_name(m.group(1))
        col_name = m.group(2).upper()
        comment = m.group(3)
        if table_name not in new_table_names:
            result["col_comments"].append({
                "table_name": table_name, "col_name": col_name, "comment": comment,
            })
    # COMMENT ON TABLE (独立)
    for m in re.finditer(rf"COMMENT\s+ON\s+TABLE\s+({SQL_NAME})\s+IS\s+'([^']*)'", sql_content, re.I):
        table_name = _table_name(m.group(1))
        if table_name not in new_table_names:
            result["table_comments"].append({
                "table_name": table_name, "comment": m.group(2),
            })
    return result


# ===================== DOCX 自动检测 =====================

DATA_TABLE_HEADERS = ["序号", "字段名称", "字段描述", "类型", "长度", "允许空", "缺省值"]
COLUMN_HEADER_ALIASES = {
    "sequence": {"序号", "编号", "序列", "#"},
    "name": {"字段名称", "字段名", "字段", "列名", "COLUMNNAME"},
    "description": {"字段描述", "字段说明", "字段含义", "中文名称", "名称", "说明"},
    "type": {"类型", "数据类型", "字段类型"},
    "length": {"长度", "数据长度", "字段长度", "精度"},
    "nullable": {"允许空", "允许为空", "是否为空", "是否可空", "可空", "能否为空"},
    "not_null": {"非空", "是否非空"},
    "primary_key": {"主键", "是否主键"},
    "default": {"缺省值", "默认值", "缺省", "默认"},
    "remarks": {"备注说明", "备注"},
}
DESCRIPTION_PREFIXES = ("描述：", "说明：", "描述:", "说明:", "Description:")


def _get_para_style_id(elem):
    """获取段落样式 ID"""
    pPr = elem.find(qn("w:pPr"))
    if pPr is not None:
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            return pStyle.get(qn("w:val"))
    return None


def _get_para_style_name(paragraph):
    """获取 python-docx 段落样式名称"""
    try:
        return paragraph.style.name
    except:
        return ""


def _get_para_text(elem):
    """获取段落 XML 元素的完整文本"""
    texts = elem.findall(".//" + qn("w:t"))
    return "".join(t.text or "" for t in texts).strip()


def _get_row_texts(row):
    """获取表格行各单元格文本"""
    return [cell.text.strip() for cell in row.cells]


def _normalize_header(text):
    return re.sub(r"[\s_\-/]", "", str(text or "")).upper()


def _column_map(table, profile=None):
    """Return semantic indexes for a field table whose header is in the first three rows."""
    if len(table.rows) == 0:
        return {}
    aliases = {
        key: {_normalize_header(item) for item in values}
        for key, values in COLUMN_HEADER_ALIASES.items()
    }
    required = {"name", "description", "type"}
    for header_row_idx, row in enumerate(table.rows[:3]):
        header_texts = _get_row_texts(row)
        result = {}
        for index, header in enumerate(header_texts):
            normalized = _normalize_header(header)
            for key, values in aliases.items():
                if key not in result and normalized in values:
                    result[key] = index
                    break
        if not required.issubset(result) or len(result) < 4:
            continue
        if profile:
            expected = {_normalize_header(item) for item in profile.get("headers", [])}
            actual = {_normalize_header(item) for item in header_texts}
            if expected and len(expected & actual) < min(4, len(expected)):
                continue
        result["_header_row_index"] = header_row_idx
        return result
    return {}


def _is_data_table(table, profile=None):
    return bool(_column_map(table, profile))


def _is_description(text):
    return any(str(text).startswith(prefix) for prefix in DESCRIPTION_PREFIXES)


def _looks_like_table_name(text):
    return bool(re.fullmatch(
        r"(?:[A-Za-z_][\w$#]*\.)?[A-Za-z_][\w$#]*(?:\s*\[[^\]]*\])?",
        text or "",
    ))


def _table_identity(text):
    """Extract a database identifier and optional Chinese description from a title."""
    text = str(text or "").strip()
    bracketed = re.fullmatch(
        r"((?:[A-Za-z_][\w$#]*\.)?[A-Za-z_][\w$#]*)\s*\[([^\]]*)\]",
        text,
    )
    if bracketed:
        return bracketed.group(1).rsplit(".", 1)[-1], bracketed.group(2).strip(), "bracket"
    exact = re.fullmatch(r"(?:[A-Za-z_][\w$#]*\.)?([A-Za-z_][\w$#]*)", text)
    if exact:
        return exact.group(1), "", "plain"
    tokens = re.findall(r"[A-Za-z_][A-Za-z0-9_$#]*", text)
    if tokens:
        identifier = tokens[-1]
        prefix = text[:text.rfind(identifier)].strip(" -_—[]（）()")
        return identifier, prefix, "embedded"
    return "", "", "plain"


def _revision_column_map(table):
    """Infer logical revision-history columns, including merged Word cells."""
    aliases = {
        "version": ("版本", "版本号"),
        "change_type": ("类型", "变更类型"),
        "date": ("日期", "修订日期"),
        "author": ("作者", "编写人", "修订人"),
        "description": ("说明", "修改内容", "修订内容", "变更内容"),
    }
    for row_index, row in enumerate(table.rows):
        result = {}
        for index, cell in enumerate(row.cells):
            normalized = re.sub(r"\s+", "", cell.text)
            for key, names in aliases.items():
                if key not in result and any(name in normalized for name in names):
                    result[key] = index
                    break
        if {"version", "date", "author", "description"}.issubset(result):
            return result, row_index
    return {}, None


def _analyze_revision_history(table):
    column_map, header_row_idx = _revision_column_map(table)
    version_rows = []
    version_index = column_map.get("version")
    for row_index, row in enumerate(table.rows):
        if not row.cells or version_index is None or version_index >= len(row.cells):
            continue
        match = re.fullmatch(r"V(\d+(?:\.\d+)*)", row.cells[version_index].text.strip(), re.I)
        if match:
            version_rows.append((row_index, tuple(int(item) for item in match.group(1).split("."))))
    last_row_idx = version_rows[-1][0] if version_rows else None
    blank_rows_after = []
    if last_row_idx is not None:
        for row_index in range(last_row_idx + 1, len(table.rows)):
            if all(not cell.text.strip() for cell in table.rows[row_index].cells):
                blank_rows_after.append(row_index)
    if column_map.get("date") == 0 and column_map.get("version") == 1 and len(table.rows[0].cells) >= 6:
        layout = "date_first6"
    elif "change_type" in column_map:
        layout = "type6"
    elif len(table.rows[0].cells) == 5 and column_map.get("date") in {1, 2}:
        layout = "date_merged5"
    elif column_map:
        layout = "standard4"
    else:
        layout = "unknown"
    return {
        "column_map": column_map,
        "header_row_idx": header_row_idx,
        "version_rows": version_rows,
        "last_version_row_idx": last_row_idx,
        "blank_rows_after": blank_rows_after,
        "layout": layout,
    }


def analyze_document(doc, profile=None):
    """
    自动分析文档结构，返回：
    {
        "data_tables": {
            "TABLE_NAME": {
                "heading_style_id": "2",      # 标题样式ID
                "heading_style_name": "Heading 2",
                "heading_para_idx": 47,        # 标题段落索引
                "table_idx": 1,                # 表格索引
                "description": "描述：xxx",     # 描述文本（可能为空）
                "description_format": {        # null 如果没有描述段落
                    "prefix": "描述：",         # 描述前缀
                    "empty_paras_before": 0,   # 标题和描述之间的空段落数
                    "empty_paras_after": 2,    # 描述和表格之间的空段落数
                },
                "empty_paras_after_table": 1,  # 表格后的空段落数
            }
        },
        "version_table_idx": 0,               # 版本历史表索引，None 如果没有
        "toc": {
            "style_id": "6",                   # TOC 样式 ID
            "style_name": "toc 2",             # TOC 样式名称
            "last_entry_idx": 237,             # 最后一条 TOC 条目段落索引
        } or None,
        "heading_levels_used": ["2"],          # 文档中用于表名的标题样式ID列表
    }
    """
    body = doc.element.body
    body_children = list(body)
    body_positions = {id(child): index for index, child in enumerate(body_children)}
    paragraph_style_names = {
        id(paragraph._element): (paragraph.style.name if paragraph.style else "")
        for paragraph in doc.paragraphs
    }

    # Find each field table and its nearest table-name paragraph. Some systems
    # (notably OCR) deliberately use Normal paragraphs instead of Heading 2.
    table_idx = 0
    para_idx = 0
    heading_seq = []
    pending_paragraphs = []

    for child in body_children:
        if child.tag == qn("w:p"):
            style_id = _get_para_style_id(child)
            text = _get_para_text(child)
            pending_paragraphs.append({
                "text": text, "elem": child, "index": para_idx, "style_id": style_id,
            })
            para_idx += 1
        elif child.tag == qn("w:tbl"):
            table = doc.tables[table_idx]
            column_map = _column_map(table, profile)
            if column_map:
                header_row_idx = column_map.get("_header_row_index", 0)
                embedded_text = ""
                embedded_name = ""
                if header_row_idx > 0:
                    embedded_text = table.rows[header_row_idx - 1].cells[0].text.strip()
                    embedded_name, _, _ = _table_identity(embedded_text)
                if embedded_name:
                    heading_seq.append(({
                        "text": embedded_text,
                        "elem": None,
                        "index": None,
                        "style_id": None,
                    }, [], table_idx, column_map))
                else:
                    non_empty = [item for item in pending_paragraphs if item["text"]]
                    candidates = [
                        item for item in non_empty
                        if not _is_description(item["text"]) and _looks_like_table_name(item["text"])
                    ]
                    if not candidates:
                        candidates = [item for item in non_empty if not _is_description(item["text"])]
                    if candidates:
                        heading = candidates[-1]
                        between = [item for item in pending_paragraphs if item["index"] > heading["index"]]
                        heading_seq.append((heading, between, table_idx, column_map))
            table_idx += 1
            pending_paragraphs = []

    # 从 heading_seq 构建 data_tables，补充段落级信息
    data_table_entries = {}
    heading_styles = Counter()

    for i, (heading, between, tidx, column_map) in enumerate(heading_seq):
        table_name, title_description, heading_format = _table_identity(heading["text"])
        htext = (table_name or heading["text"]).upper()
        helem = heading["elem"]
        hidx = heading["index"]
        hstyle = heading["style_id"]
        heading_styles[hstyle] += 1

        # 通过 python-docx 段落 API 获取样式名称
        hstyle_name = paragraph_style_names.get(id(helem), "")

        # 找标题后的描述段落
        description = title_description
        desc_format = None
        empty_after_table = 1  # 默认

        desc_index = None
        for index, item in enumerate(between):
            if _is_description(item["text"]):
                desc_index = index
                desc_text = item["text"]
                description = desc_text
                prefix = next(prefix for prefix in DESCRIPTION_PREFIXES if desc_text.startswith(prefix))
                desc_format = {
                    "prefix": prefix,
                    "empty_paras_before": sum(not p["text"] for p in between[:index]),
                }
                break
        empty_after_desc = 0
        if desc_index is not None:
            empty_after_desc = sum(not p["text"] for p in between[desc_index + 1:])

        if desc_format:
            desc_format["empty_paras_after"] = empty_after_desc

        # 找表后的空段落数
        # 简单策略：找表后面的空段落
        ea = 0
        table_position = body_positions.get(id(doc.tables[tidx]._element), -1)
        for child in body_children[table_position + 1:]:
            if child.tag == qn("w:tbl"):
                break
            if child.tag == qn("w:p"):
                if not _get_para_text(child):
                    ea += 1
                else:
                    break
        empty_after_table = max(1, ea)

        data_table_entries[htext] = {
            "heading_style_id": hstyle,
            "heading_style_name": hstyle_name,
            "heading_para_idx": hidx,
            "table_idx": tidx,
            "description": description,
            "description_format": desc_format,
            "heading_format": heading_format,
            "embedded_title_row_idx": (
                column_map.get("_header_row_index", 0) - 1
                if column_map.get("_header_row_index", 0) > 0 else None
            ),
            "empty_paras_after_table": empty_after_table,
            "heading_elem": helem,
            "column_map": column_map,
        }

    # 找出最常用的标题样式
    most_common_heading = heading_styles.most_common(1)
    default_heading_style = most_common_heading[0][0] if most_common_heading else None
    default_heading_name = ""
    for v in data_table_entries.values():
        if v["heading_style_id"] == default_heading_style:
            default_heading_name = v["heading_style_name"]
            break
    default_heading_elem = next(
        (v["heading_elem"] for v in data_table_entries.values()
         if v["heading_style_id"] == default_heading_style),
        next((v["heading_elem"] for v in data_table_entries.values()), None),
    )

    # 版本历史表检测：按语义列识别，兼容“日期在前、版本在后”的设计文档。
    version_table_idx = None
    for ti, table in enumerate(doc.tables):
        if ti in [v["table_idx"] for v in data_table_entries.values()]:
            continue  # 跳过数据表
        revision = _analyze_revision_history(table)
        if revision["version_rows"] and {
            "version", "date", "author", "description"
        }.issubset(revision["column_map"]):
            version_table_idx = ti
            break

    # TOC 检测
    toc_info = None
    last_toc_elem = None
    last_toc_idx = None
    toc_style_id = None
    toc_style_name = None
    for i, p in enumerate(doc.paragraphs):
        sn = p.style.name if p.style else ""
        if "toc" in sn.lower() or "TOC" in sn or "目录" in sn:
            toc_style_id = _get_para_style_id(p._element)
            toc_style_name = sn
            last_toc_elem = p._element
            last_toc_idx = i
    if last_toc_elem is not None:
        toc_info = {
            "style_id": toc_style_id,
            "style_name": toc_style_name,
            "last_entry_elem": last_toc_elem,
            "last_entry_idx": last_toc_idx,
        }

    # 提取典型描述格式（用于新表）
    typical_desc_format = None
    for v in data_table_entries.values():
        if v["description_format"]:
            typical_desc_format = v["description_format"]
            break
    heading_formats = Counter(v.get("heading_format", "plain") for v in data_table_entries.values())
    typical_heading_format = heading_formats.most_common(1)[0][0] if heading_formats else "plain"

    # Pick the most common table geometry, not simply the last table. This
    # avoids propagating formatting defects introduced by an older updater.
    geometry_counts = Counter()
    geometry_first_table = {}
    for value in data_table_entries.values():
        table = doc.tables[value["table_idx"]]
        signature = tuple(
            cell._tc.tcPr.tcW.get(qn("w:w")) if cell._tc.tcPr.tcW is not None else ""
            for cell in table.rows[0].cells
        )
        geometry_counts[signature] += 1
        geometry_first_table.setdefault(signature, value["table_idx"])
    common_geometry = geometry_counts.most_common(1)
    template_table_idx = geometry_first_table[common_geometry[0][0]] if common_geometry else None

    return {
        "data_tables": data_table_entries,
        "version_table_idx": version_table_idx,
        "revision_history": (
            _analyze_revision_history(doc.tables[version_table_idx])
            if version_table_idx is not None else {"layout": "none", "column_map": {}}
        ),
        "toc": toc_info,
        "default_heading_style_id": default_heading_style,
        "default_heading_style_name": default_heading_name,
        "default_heading_elem": default_heading_elem,
        "typical_desc_format": typical_desc_format,
        "typical_heading_format": typical_heading_format,
        "template_table_idx": template_table_idx,
        "total_tables": len(doc.tables),
        "profile": profile,
    }


def detect_existing_changes(docx_path, sql_content):
    """Preflight supported SQL against the current document.

    Each returned item is already represented by a table, field, type, or
    identical comment in the DOCX and can be shown to the user for confirmation.
    """
    doc = Document(docx_path)
    profile = match_document_template(docx_path)
    analysis = analyze_document(doc, profile)
    data_tables = analysis["data_tables"]
    parsed = parse_sql(sql_content)
    found = []

    for item in parsed["new_tables"]:
        if item["table_name"] in data_tables:
            found.append({"kind": "table", "table": item["table_name"], "field": "", "detail": "表已存在"})

    for item in parsed["alter_adds"]:
        info = data_tables.get(item["table_name"])
        if not info:
            continue
        table = doc.tables[info["table_idx"]]
        for column in item["columns"]:
            if find_col_row(table, column["name"], info["column_map"]) is not None:
                found.append({"kind": "field", "table": item["table_name"], "field": column["name"], "detail": "字段已存在"})

    for item in parsed["alter_modifies"]:
        info = data_tables.get(item["table_name"])
        if not info:
            continue
        table = doc.tables[info["table_idx"]]
        for column in item["columns"]:
            column_map = info["column_map"]
            row_index = find_col_row(table, column["name"], column_map)
            if row_index is None:
                continue
            row = table.rows[row_index]
            current = (
                row.cells[column_map["type"]].text.strip().upper(),
                row.cells[column_map["length"]].text.strip() if "length" in column_map else "",
            )
            requested = (column.get("type", "").upper(), column.get("length", ""))
            if current == requested:
                found.append({"kind": "modify", "table": item["table_name"], "field": column["name"], "detail": "字段类型和长度已一致"})

    for item in parsed["col_comments"]:
        info = data_tables.get(item["table_name"])
        if not info:
            continue
        table = doc.tables[info["table_idx"]]
        column_map = info["column_map"]
        row_index = find_col_row(table, item["col_name"], column_map)
        if row_index is not None and table.rows[row_index].cells[column_map["description"]].text.strip() == item["comment"].strip():
            found.append({"kind": "comment", "table": item["table_name"], "field": item["col_name"], "detail": "字段说明已一致"})

    for item in parsed["table_comments"]:
        info = data_tables.get(item["table_name"])
        if info and info["description"].strip() in {
            item["comment"].strip(), f"描述：{item['comment'].strip()}", f"描述:{item['comment'].strip()}"
        }:
            found.append({"kind": "comment", "table": item["table_name"], "field": "", "detail": "表说明已一致"})
    return found


# ===================== DOCX 写入操作 =====================

def set_cell_text(cell, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size_emu=127000):
    """设置表格单元格"""
    p = cell.paragraphs[0]
    p.alignment = alignment
    for r in p.runs:
        r.text = ""
    run = p.add_run(str(text))
    run.font.size = font_size_emu
    run.font.bold = bold


def find_col_row(table, col_name, column_map=None):
    """在表格中查找字段行"""
    column_map = column_map or _column_map(table)
    name_index = column_map.get("name")
    if name_index is None:
        return None
    start_row = column_map.get("_header_row_index", 0) + 1
    for ri in range(start_row, len(table.rows)):
        if name_index < len(table.rows[ri].cells) and table.rows[ri].cells[name_index].text.strip().upper() == col_name.upper():
            return ri
    return None


def _append_formatted_row(table):
    """Append a row cloned from the table's existing data-row formatting."""
    source_row = table.rows[-1] if len(table.rows) > 1 else table.rows[0]
    table._tbl.append(deepcopy(source_row._tr))
    new_row = table.rows[-1]
    for cell in new_row.cells:
        set_cell_text(cell, "")
    return new_row


def _write_column_row(row, column_map, column, sequence):
    values = {
        "sequence": sequence,
        "name": column["name"],
        "description": column.get("comment", ""),
        "type": column.get("type", ""),
        "length": column.get("length", ""),
        "nullable": "√" if column.get("nullable", True) else "",
        "not_null": "" if column.get("nullable", True) else "√",
        "primary_key": "√" if column.get("primary_key", False) else "",
        "default": column.get("default", ""),
        "remarks": "",
    }
    for key, value in values.items():
        index = column_map.get(key)
        if index is not None and index < len(row.cells):
            set_cell_text(row.cells[index], value)


def add_column_rows_to_table(table, columns, column_map=None):
    """向表格追加字段行并重新编号"""
    column_map = column_map or _column_map(table)
    for col in columns:
        new_row = _append_formatted_row(table)
        header_row_idx = column_map.get("_header_row_index", 0)
        seq = str(len(table.rows) - header_row_idx - 1)
        _write_column_row(new_row, column_map, col, seq)
    _renumber_table(table, column_map)


def _renumber_table(table, column_map=None):
    """重新编号表格序号列"""
    column_map = column_map or _column_map(table)
    sequence_index = column_map.get("sequence")
    if sequence_index is None:
        return
    start_row = column_map.get("_header_row_index", 0) + 1
    for sequence, ri in enumerate(range(start_row, len(table.rows)), 1):
        set_cell_text(table.rows[ri].cells[sequence_index], str(sequence))


def modify_column_in_table(table, col_name, new_col, column_map=None):
    """修改表格中字段的类型/长度"""
    column_map = column_map or _column_map(table)
    ri = find_col_row(table, col_name, column_map)
    if ri is None:
        return False
    row = table.rows[ri]
    if new_col.get("type"):
        set_cell_text(row.cells[column_map["type"]], new_col["type"])
    if new_col.get("length") and "length" in column_map:
        set_cell_text(row.cells[column_map["length"]], new_col["length"])
    return True


def update_column_comment_in_table(table, col_name, comment, column_map=None):
    """更新字段描述"""
    column_map = column_map or _column_map(table)
    ri = find_col_row(table, col_name, column_map)
    if ri is None:
        return False
    set_cell_text(table.rows[ri].cells[column_map["description"]], comment)
    return True


def update_table_description(doc, analysis, table_name, new_desc):
    """更新标题后的描述段落"""
    dt = analysis["data_tables"].get(table_name.upper())
    if not dt:
        return False
    if dt.get("heading_format") == "bracket" and dt.get("heading_elem") is not None:
        text_nodes = dt["heading_elem"].findall(".//" + qn("w:t"))
        if text_nodes:
            text_nodes[0].text = f"{table_name.upper()} [{new_desc}]"
            for node in text_nodes[1:]:
                node.text = ""
            return True
    title_row_idx = dt.get("embedded_title_row_idx")
    if title_row_idx is not None:
        row = doc.tables[dt["table_idx"]].rows[title_row_idx]
        seen_cells = set()
        for cell in row.cells:
            cell_key = id(cell._tc)
            if cell_key not in seen_cells:
                set_cell_text(cell, f"{new_desc} {table_name.upper()}".strip())
                seen_cells.add(cell_key)
        return True
    if not dt["description_format"]:
        return False
    helem = dt["heading_elem"]
    body = doc.element.body
    found = False
    for child in body:
        if child is helem:
            found = True
            continue
        if found and child.tag == qn("w:p"):
            text = _get_para_text(child)
            prefix = dt["description_format"]["prefix"]
            if text.startswith(prefix):
                for t in child.findall(".//" + qn("w:t")):
                    if t.text and t.text.startswith(prefix):
                        t.text = f"{prefix}{new_desc}"
                        return True
            break
    return False


def add_new_table_section(doc, analysis, table_name, description, columns):
    """向文档末尾添加新表章节，格式克隆自已有数据表"""
    dt = analysis["data_tables"]
    tf = analysis["typical_desc_format"]
    hs_name = analysis["default_heading_style_name"]
    heading_format = analysis.get("typical_heading_format", "plain")

    # 取最后一个数据表的格式作为模板
    last_dt = None
    for v in dt.values():
        last_dt = v

    # 空行间隔
    trailing_empties = last_dt["empty_paras_after_table"] if last_dt else 1
    for _ in range(trailing_empties):
        doc.add_paragraph("")

    if heading_format != "embedded":
        heading_text = (
            f"{table_name.upper()} [{description}]"
            if heading_format == "bracket" and description else table_name.upper()
        )
        # 标题：有标题样式时直接沿用；OCR 等 Normal 标题克隆原段落格式。
        if hs_name:
            doc.add_paragraph(heading_text, style=hs_name)
        elif analysis.get("default_heading_elem") is not None:
            heading_elem = deepcopy(analysis["default_heading_elem"])
            text_nodes = heading_elem.findall(".//" + qn("w:t"))
            if text_nodes:
                text_nodes[0].text = heading_text
                for node in text_nodes[1:]:
                    node.text = ""
            doc.element.body.insert(-1, heading_elem)
        else:
            doc.add_paragraph(heading_text)

        # 描述段落
        if tf and description:
            for _ in range(tf.get("empty_paras_before", 0)):
                doc.add_paragraph("")
            doc.add_paragraph(f"{tf['prefix']}{description}")
            for _ in range(tf.get("empty_paras_after", 2)):
                doc.add_paragraph("")
        elif heading_format != "bracket":
            doc.add_paragraph("")
            doc.add_paragraph("")

    # Clone the document's dominant table template so column widths, borders,
    # cell margins and paragraph formatting stay consistent with the source.
    template_idx = analysis.get("template_table_idx")
    if template_idx is not None and len(doc.tables[template_idx].rows) > 1:
        template = doc.tables[template_idx]
        template_map = _column_map(template, analysis.get("profile"))
        header_row_idx = template_map.get("_header_row_index", 0)
        table_element = deepcopy(template._tbl)
        while len(table_element.tr_lst) > header_row_idx + 2:
            table_element.remove(table_element.tr_lst[-1])
        doc.element.body.insert(-1, table_element)
        table = doc.tables[-1]
    else:
        headers = (analysis.get("profile") or {}).get("headers") or DATA_TABLE_HEADERS
        table = doc.add_table(rows=2, cols=len(headers))
        table.style = "Table Grid"
        for ci, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[ci], header, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    column_map = _column_map(table, analysis.get("profile"))
    if not column_map:
        raise ValueError(f"模板字段表表头不兼容，无法安全新增表 {table_name}")

    if not columns and len(table.rows) > 1:
        data_row_idx = column_map.get("_header_row_index", 0) + 1
        if data_row_idx < len(table.rows):
            table._tbl.remove(table.rows[data_row_idx]._tr)
        return table

    if heading_format == "embedded" and column_map.get("_header_row_index", 0) > 0:
        title_row = table.rows[column_map["_header_row_index"] - 1]
        seen_cells = set()
        for cell in title_row.cells:
            cell_key = id(cell._tc)
            if cell_key not in seen_cells:
                set_cell_text(cell, f"{description} {table_name.upper()}".strip())
                seen_cells.add(cell_key)

    for ri, col in enumerate(columns):
        data_row_idx = column_map.get("_header_row_index", 0) + 1
        row = table.rows[data_row_idx] if ri == 0 else _append_formatted_row(table)
        _write_column_row(row, column_map, col, str(ri + 1))
    return table


def add_toc_entry(doc, analysis, table_name):
    """添加 TOC 条目"""
    toc = analysis["toc"]
    if not toc:
        return False
    body = doc.element.body
    p = doc.add_paragraph(table_name.upper(), style=toc["style_name"])
    body.remove(p._element)
    last_elem = toc["last_entry_elem"]
    next_elem = last_elem.getnext()
    if next_elem is not None:
        next_elem.addprevious(p._element)
    else:
        body.append(p._element)
    toc["last_entry_elem"] = p._element
    return True


def _set_cell_text_preserve_format(cell, text):
    """Replace text while retaining the cloned revision row's run formatting."""
    paragraphs = list(cell.paragraphs)
    first_paragraph = paragraphs[0]
    first_run = first_paragraph.runs[0] if first_paragraph.runs else first_paragraph.add_run()
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.text = ""
    for paragraph in paragraphs[1:]:
        cell._tc.remove(paragraph._element)
    first_run.text = str(text)


def _revision_change_type(changes):
    has_add = any("新增" in item or "增加" in item for item in changes)
    has_update = any("修改" in item or "更新" in item for item in changes)
    if has_add and has_update:
        return "A/U"
    if has_add:
        return "A"
    return "U"


def _prepare_revision_row(table, revision):
    """Clone the last real revision row and consume a reserved blank row."""
    source_idx = revision["last_version_row_idx"]
    if source_idx is None:
        raise ValueError("修订历史表未找到可继承格式的版本记录行")
    cloned_tr = deepcopy(table.rows[source_idx]._tr)
    blank_rows = revision.get("blank_rows_after") or []
    if blank_rows:
        target_idx = blank_rows[0]
        blank_tr = table.rows[target_idx]._tr
        blank_tr.addprevious(cloned_tr)
        table._tbl.remove(blank_tr)
    else:
        target_idx = len(table.rows)
        table._tbl.append(cloned_tr)
    target_row = table.rows[target_idx]
    seen_cells = set()
    for cell in target_row.cells:
        cell_key = id(cell._tc)
        if cell_key not in seen_cells:
            _set_cell_text_preserve_format(cell, "")
            seen_cells.add(cell_key)
    return target_row


def update_version_history(doc, analysis, author, changes, update_date=None, revision_type=None):
    """更新版本历史"""
    vt_idx = analysis["version_table_idx"]
    today = update_date or date.today().strftime("%Y-%m-%d")

    if vt_idx is not None:
        table = doc.tables[vt_idx]
        revision = _analyze_revision_history(table)
        column_map = revision["column_map"]
        required = {"version", "date", "author", "description"}
        if not required.issubset(column_map):
            raise ValueError("修订历史表列无法识别，已停止写入以避免版本记录错列")
        versions = [parts for _, parts in revision["version_rows"]]
        parts = list(max(versions) if versions else (1, 0, 0))
        parts[-1] += 1
        new_ver = "V" + ".".join(str(item) for item in parts)
        desc = "\n".join(changes)
        new_row = _prepare_revision_row(table, revision)
        values = {
            "version": new_ver,
            "date": today,
            "author": author,
            "description": desc,
            "change_type": revision_type or _revision_change_type(changes),
        }
        written_cells = set()
        row_cells = new_row.cells
        for key, value in values.items():
            index = column_map.get(key)
            if index is None:
                continue
            cell = row_cells[index]
            cell_key = id(cell._tc)
            if cell_key not in written_cells:
                _set_cell_text_preserve_format(cell, value)
                written_cells.add(cell_key)
        return new_ver
    else:
        # 无版本历史表，追加到文档末尾
        doc.add_paragraph("")
        doc.add_paragraph("─" * 40)
        doc.add_paragraph(f"自动更新记录 - {today} by {author}")
        for ch in changes:
            doc.add_paragraph(f"  • {ch}")
        return f"auto-{today}"


# ===================== 核心处理 =====================
def process(docx_path, sql_path, author="System", output_path=None, backup=True, update_date=None):
    profile = match_document_template(docx_path)
    report = {
        "changes": [], "skipped": [], "save_path": None, "version": None,
        "template": profile,
    }
    # ── 读取 SQL ──
    if not os.path.exists(sql_path):
        raise FileNotFoundError(f"SQL 文件不存在: {sql_path}")
    enc = detect_encoding(sql_path)
    print(f"[信息] SQL 文件编码: {enc}")
    with open(sql_path, "r", encoding=enc) as f:
        sql_content = f.read()

    # ── 解析 SQL ──
    parsed = parse_sql(sql_content)
    total = (len(parsed["new_tables"]) + len(parsed["alter_adds"]) +
             len(parsed["alter_modifies"]) + len(parsed["col_comments"]) +
             len(parsed["table_comments"]))
    if total == 0:
        print("[警告] 未发现任何有效的 DDL/注释语句")
        return report

    print(f"[信息] 解析到 {total} 项变更:")
    for t in parsed["new_tables"]:
        print(f"  + 新增表 {t['table_name']} ({len(t['columns'])} 列): {t['table_comment']}")
    for a in parsed["alter_adds"]:
        print(f"  + ALTER ADD  {a['table_name']}: {', '.join(c['name'] for c in a['columns'])}")
    for m in parsed["alter_modifies"]:
        print(f"  ~ ALTER MOD  {m['table_name']}: {', '.join(c['name'] for c in m['columns'])}")
    for cc in parsed["col_comments"]:
        print(f"  * COMMENT    {cc['table_name']}.{cc['col_name']}")
    for tc in parsed["table_comments"]:
        print(f"  * TABLE CMT  {tc['table_name']}")

    # ── 读取 DOCX ──
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")
    if backup and output_path is None:
        backup_path = docx_path + ".bak"
        shutil.copy2(docx_path, backup_path)
        print(f"[信息] 已备份: {backup_path}")

    doc = Document(docx_path)

    # ── 自动分析 ──
    analysis = analyze_document(doc, profile)
    data_tables = analysis["data_tables"]
    existing_names = set(data_tables.keys())
    revision_layout = analysis["revision_history"]["layout"]
    report["revision_layout"] = revision_layout
    if profile and profile.get("revision_layout") != revision_layout:
        raise ValueError(
            f"{profile['system']} 修订历史结构与最新模板不一致："
            f"期望 {profile.get('revision_layout')}，实际 {revision_layout}。已停止写入。"
        )

    print(f"[分析] 检测到 {len(data_tables)} 个数据表")
    if profile:
        print(f"[模板] {profile['system']} → {profile['template']} ({profile['confidence']})")
    else:
        print("[模板] 文件名未匹配内置系统，使用文档自身结构进行兼容分析")
    print(f"[分析] 标题样式: {analysis['default_heading_style_name']} (id={analysis['default_heading_style_id']})")
    print(f"[分析] 版本历史表: {'Table ' + str(analysis['version_table_idx']) if analysis['version_table_idx'] is not None else '未检测到'}")
    print(f"[分析] 修订历史布局: {revision_layout}")
    print(f"[分析] TOC 目录: {'已检测到' if analysis['toc'] else '未检测到'}")
    if analysis["typical_desc_format"]:
        print(f"[分析] 描述格式: '{analysis['typical_desc_format']['prefix']}...'")

    # ── 执行变更 ──
    changes = []
    revision_types = set()

    # 新增表
    new_t = [t for t in parsed["new_tables"] if t["table_name"] not in existing_names]
    skipped_t = [t for t in parsed["new_tables"] if t["table_name"] in existing_names]
    for t in skipped_t:
        print(f"[跳过] 表已存在: {t['table_name']}")
        report["skipped"].append(f"表已存在: {t['table_name']}")
    for t in new_t:
        add_toc_entry(doc, analysis, t["table_name"])
        add_new_table_section(doc, analysis, t["table_name"], t["table_comment"], t["columns"])
        changes.append(f"新增{t['table_name'].lower()}表（{t['table_comment']}）")
        revision_types.add("A")
        existing_names.add(t["table_name"])
        print(f"[新增表] {t['table_name']} ({len(t['columns'])} 列)")

    # New tables are now part of the document; refresh indexes before ALTERs.
    analysis = analyze_document(doc, profile)
    data_tables = analysis["data_tables"]

    # ALTER ADD
    alter_add_cols = set()
    for a in parsed["alter_adds"]:
        if a["table_name"] not in data_tables:
            print(f"[警告] 目标表不存在: {a['table_name']}，跳过 ALTER ADD")
            continue
        info = data_tables[a["table_name"]]
        table = doc.tables[info["table_idx"]]
        column_map = info["column_map"]
        missing_columns = [c for c in a["columns"] if find_col_row(table, c["name"], column_map) is None]
        duplicate_columns = [c for c in a["columns"] if find_col_row(table, c["name"], column_map) is not None]
        for column in duplicate_columns:
            message = f"字段已存在: {a['table_name']}.{column['name']}"
            print(f"[跳过] {message}")
            report["skipped"].append(message)
        if not missing_columns:
            continue
        add_column_rows_to_table(table, missing_columns, column_map)
        col_names = "、".join(c["name"] for c in missing_columns)
        changes.append(f"{a['table_name']}表增加{col_names}字段")
        revision_types.add("A")
        print(f"[ALTER ADD] {a['table_name']}: 新增 {len(missing_columns)} 个字段")
        for c in missing_columns:
            alter_add_cols.add((a["table_name"].upper(), c["name"].upper()))

    # ALTER MODIFY
    for m in parsed["alter_modifies"]:
        if m["table_name"] not in data_tables:
            print(f"[警告] 目标表不存在: {m['table_name']}，跳过 ALTER MODIFY")
            continue
        info = data_tables[m["table_name"]]
        table = doc.tables[info["table_idx"]]
        column_map = info["column_map"]
        for col in m["columns"]:
            row_index = find_col_row(table, col["name"], column_map)
            if row_index is not None:
                row = table.rows[row_index]
                current = (
                    row.cells[column_map["type"]].text.strip().upper(),
                    row.cells[column_map["length"]].text.strip() if "length" in column_map else "",
                )
                requested = (col.get("type", "").upper(), col.get("length", ""))
                if current == requested:
                    message = f"字段类型和长度已一致: {m['table_name']}.{col['name']}"
                    print(f"[跳过] {message}")
                    report["skipped"].append(message)
                    continue
            ok = modify_column_in_table(table, col["name"], col, column_map)
            if ok:
                changes.append(
                    f"{m['table_name']}表修改{col['name']}字段类型为{col['type']}"
                    + (f"({col['length']})" if col["length"] else "")
                )
                revision_types.add("U")
                print(f"[ALTER MOD] {m['table_name']}.{col['name']} → {col['type']}({col['length']})")
            else:
                print(f"[警告] 字段不存在: {m['table_name']}.{col['name']}")

    # COMMENT ON COLUMN（跳过 alter_adds 已处理的）
    col_comments = [
        cc for cc in parsed["col_comments"]
        if (cc["table_name"].upper(), cc["col_name"].upper()) not in alter_add_cols
    ]
    for cc in col_comments:
        if cc["table_name"] not in data_tables:
            print(f"[警告] 目标表不存在: {cc['table_name']}，跳过 COMMENT")
            continue
        info = data_tables[cc["table_name"]]
        table = doc.tables[info["table_idx"]]
        column_map = info["column_map"]
        row_index = find_col_row(table, cc["col_name"], column_map)
        if row_index is not None and table.rows[row_index].cells[column_map["description"]].text.strip() == cc["comment"].strip():
            message = f"字段说明已一致: {cc['table_name']}.{cc['col_name']}"
            print(f"[跳过] {message}")
            report["skipped"].append(message)
            continue
        ok = update_column_comment_in_table(table, cc["col_name"], cc["comment"], column_map)
        if ok:
            changes.append(f"{cc['table_name']}表更新{cc['col_name']}字段注释")
            revision_types.add("U")
            print(f"[COMMENT] {cc['table_name']}.{cc['col_name']} → {cc['comment']}")
        else:
            print(f"[警告] 字段不存在: {cc['table_name']}.{cc['col_name']}")

    # COMMENT ON TABLE
    for tc in parsed["table_comments"]:
        if tc["table_name"] not in data_tables:
            print(f"[警告] 目标表不存在: {tc['table_name']}，跳过 TABLE CMT")
            continue
        current_description = data_tables[tc["table_name"]]["description"].strip()
        if current_description in {tc["comment"].strip(), f"描述：{tc['comment'].strip()}", f"描述:{tc['comment'].strip()}"}:
            message = f"表说明已一致: {tc['table_name']}"
            print(f"[跳过] {message}")
            report["skipped"].append(message)
            continue
        ok = update_table_description(doc, analysis, tc["table_name"], tc["comment"])
        if ok:
            changes.append(f"更新{tc['table_name']}表描述")
            revision_types.add("U")
            print(f"[TABLE CMT] {tc['table_name']}")
        else:
            print(f"[警告] 无法更新表描述: {tc['table_name']}")

    # ── 版本历史 ──
    if changes:
        revision_type = "/".join(item for item in ("A", "U") if item in revision_types)
        new_ver = update_version_history(
            doc, analysis, author, changes, update_date,
            revision_type=revision_type or None,
        )
        report["version"] = new_ver
        print(f"[版本] {new_ver}")
    else:
        print("[结果] 没有需要更新的内容")
        save_path = output_path if output_path else docx_path
        if os.path.abspath(save_path) != os.path.abspath(docx_path):
            shutil.copy2(docx_path, save_path)
            print(f"[完成] 已保留原文档并生成整理文档: {save_path}")
        report["save_path"] = save_path
        return report

    # ── 保存 ──
    save_path = output_path if output_path else docx_path
    doc.save(save_path)
    report["changes"] = list(changes)
    report["save_path"] = save_path
    print(f"\n[完成] 文档已保存: {save_path}")
    print(f"        版本 {new_ver}，共 {len(changes)} 项变更")
    return report


# ===================== GUI =====================
def gui_mode():
    try:
        import tkinter as tk
        from tkinter import filedialog, messagebox, ttk
    except ImportError:
        print("错误：图形界面需要 tkinter 支持")
        sys.exit(1)

    root = tk.Tk()
    root.title("数据库表结构文档更新工具 v3.0")
    root.geometry("680x500")
    root.resizable(False, False)

    docx_var = tk.StringVar()
    sql_var = tk.StringVar()
    author_var = tk.StringVar(value="System")

    title = tk.Label(root, text="数据库表结构文档更新工具 v3.0",
                     font=("Microsoft YaHei", 14, "bold"))
    title.pack(pady=15)
    desc = tk.Label(root, text="通用版 — 自动检测文档格式，兼容任意系统",
                    font=("Microsoft YaHei", 9), fg="gray")
    desc.pack()

    frame = ttk.LabelFrame(root, text="文件选择", padding=10)
    frame.pack(fill="x", padx=20, pady=10)

    ttk.Label(frame, text="DOCX 文档：").grid(row=0, column=0, sticky="w", pady=5)
    ttk.Entry(frame, textvariable=docx_var, width=50).grid(row=0, column=1, padx=5)
    ttk.Button(frame, text="浏览...", command=lambda: docx_var.set(
        filedialog.askopenfilename(title="选择 DOCX", filetypes=[("Word 文档", "*.docx")])
    )).grid(row=0, column=2)

    ttk.Label(frame, text="SQL 文件：").grid(row=1, column=0, sticky="w", pady=5)
    ttk.Entry(frame, textvariable=sql_var, width=50).grid(row=1, column=1, padx=5)
    ttk.Button(frame, text="浏览...", command=lambda: sql_var.set(
        filedialog.askopenfilename(title="选择 SQL", filetypes=[("SQL 文件", "*.sql"), ("所有", "*.*")])
    )).grid(row=1, column=2)

    ttk.Label(frame, text="作者：").grid(row=2, column=0, sticky="w", pady=5)
    ttk.Entry(frame, textvariable=author_var, width=30).grid(row=2, column=1, sticky="w", padx=5)

    log_frame = ttk.LabelFrame(root, text="运行日志", padding=10)
    log_frame.pack(fill="both", expand=True, padx=20, pady=10)
    log_text = tk.Text(log_frame, height=10, wrap="word", state="disabled", font=("Consolas", 9))
    log_text.pack(fill="both", expand=True)

    def log(msg):
        log_text.config(state="normal")
        log_text.insert("end", msg + "\n")
        log_text.see("end")
        log_text.config(state="disabled")
        root.update()

    def run():
        docx = docx_var.get().strip()
        sql = sql_var.get().strip()
        author = author_var.get().strip() or "System"
        if not docx:
            messagebox.showerror("错误", "请选择 DOCX"); return
        if not sql:
            messagebox.showerror("错误", "请选择 SQL"); return
        log_text.config(state="normal"); log_text.delete("1.0", "end")
        log_text.config(state="disabled")
        try:
            import io
            old = sys.stdout
            sys.stdout = io.StringIO()
            process(docx, sql, author=author)
            output = sys.stdout.getvalue()
            sys.stdout = old
            log(output)
            log("─" * 50)
            log("✅ 完成！")
            messagebox.showinfo("完成", "文档更新成功！")
        except Exception as e:
            sys.stdout = old
            import traceback
            log(f"❌ 错误: {e}")
            log(traceback.format_exc())
            messagebox.showerror("错误", str(e))

    btn_frame = tk.Frame(root)
    btn_frame.pack(pady=10)
    ttk.Button(btn_frame, text="开始更新", command=run, width=20).pack(side="left", padx=5)
    ttk.Button(btn_frame, text="退出", command=root.destroy, width=10).pack(side="left", padx=5)

    root.mainloop()


# ===================== 入口 =====================
def main():
    if len(sys.argv) == 1:
        print("启动图形界面...")
        gui_mode()
        return
    parser = argparse.ArgumentParser(description="数据库表结构文档更新工具 v3.0 通用版")
    parser.add_argument("docx", help="DOCX 文档路径")
    parser.add_argument("sql", help="SQL 文件路径")
    parser.add_argument("--author", "-a", default="System", help="版本历史作者")
    parser.add_argument("--output", "-o", default=None, help="输出路径（默认覆盖原文件）")
    parser.add_argument("--no-backup", action="store_true", help="不备份原文件")
    args = parser.parse_args()
    try:
        process(args.docx, args.sql, author=args.author,
                output_path=args.output, backup=not args.no_backup)
    except Exception as e:
        print(f"\n错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
