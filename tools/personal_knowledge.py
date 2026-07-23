# -*- coding: utf-8 -*-
import datetime
import html
from html.parser import HTMLParser
import hashlib
import json
import os
import re
import sys
import io
import uuid
import zipfile

from docx import Document

from config import PRIVATE_KNOWLEDGE_FILE, ensure_config_dir


CATEGORIES = {
    'all': '全部内容',
    'auto_business': '车险业务',
    'interface': '接口与联调',
    'database': '数据库资料',
    'server': '服务器与日志',
    'access': '账号与访问',
    'sql_code': 'SQL 与代码',
    'learning': '学习笔记',
    'other': '其他',
}

CATEGORY_KEYWORDS = {
    'auto_business': ('车险', '投保', '保单', '保费', '险种', '承保', '批改', '理赔', '车辆', 'vin', '车架号'),
    'interface': ('接口', 'http://', 'https://', 'url', 'api', '网关', 'gateway', 'esb', 'nacos', '联调', 'curl'),
    'database': ('jdbc:', 'oracle', '数据库', 'service_name', 'sid', 'tnsnames', 'rac', '查询库', '模拟库', '生产库'),
    'server': ('服务器', '日志', 'weblogic', 'applog', '/home/', 'linux', 'docker', 'kubernetes', '主机', '目录', '路径'),
    'access': ('vpn', '堡垒机', '安全传输', '登录人', '远程设备', '账号', '用户名', '密码'),
    'sql_code': ('select ', 'insert ', 'update ', 'delete ', 'create ', 'alter ', ' from ', ' where ', 'java', 'xml', '代码'),
    'learning': ('学习', '笔记', '说明', '流程', '步骤', '问题', '解决', '注意', '命令', '日报'),
}

SENSITIVE_PATTERN = re.compile(
    r'(?i)密码|password|passwd|pwd|vpn|堡垒机|jdbc:oracle|username|用户名|账号'
)


def _resource_path(filename):
    base = getattr(sys, '_MEIPASS', os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    return os.path.join(base, 'resources', filename)


def read_text_file(path):
    with open(path, 'rb') as stream:
        data = stream.read()
    for encoding in ('utf-8-sig', 'utf-8', 'gb18030'):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode('utf-8', errors='replace')


def extract_document_text(path):
    extension = os.path.splitext(path)[1].lower()
    if extension == '.docx':
        document = Document(path)
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    lines.append('\t'.join(values))
        return '\n'.join(lines)
    if extension == '.xlsx':
        return extract_workbook_text(path)
    if extension in ('.txt', '.md', '.log', '.sql', '.json', '.xml', '.yaml', '.yml', '.csv'):
        return read_text_file(path)
    raise ValueError('暂不支持该文档格式')


def file_type_for_path(path):
    extension = os.path.splitext(str(path or ''))[1].lower()
    return {
        '.xlsx': 'EXCEL', '.docx': 'WORD', '.txt': 'TXT', '.md': 'MARKDOWN',
        '.sql': 'SQL', '.json': 'JSON', '.xml': 'XML', '.yaml': 'YAML',
        '.yml': 'YAML', '.csv': 'CSV', '.log': 'LOG',
    }.get(extension, extension.lstrip('.').upper() or '文本')


def extract_word_entry(path, builtin=False, source=None):
    document = Document(path)
    blocks = []
    plain = []
    for paragraph in document.paragraphs:
        text = paragraph.text
        plain.append(text)
        blocks.append(f'<p>{html.escape(text)}</p>')
    for table in document.tables:
        html_rows = []
        for row in table.rows:
            values = [cell.text for cell in row.cells]
            plain.append('\t'.join(values))
            html_rows.append('<tr>' + ''.join(f'<td>{html.escape(value)}</td>' for value in values) + '</tr>')
        blocks.append('<table border="1" cellspacing="0" cellpadding="4">' + ''.join(html_rows) + '</table>')
    content = '\n'.join(plain).strip()
    category = classify_content(content)
    now = datetime.datetime.now().isoformat(timespec='seconds')
    source_name = source or os.path.basename(path)
    digest = hashlib.sha256((source_name + content).encode('utf-8')).hexdigest()[:20]
    return {
        'id': f'seed-word-{digest}' if builtin else uuid.uuid4().hex,
        'title': os.path.splitext(source_name)[0], 'category': category,
        'content_type': 'word_document', 'file_type': 'WORD',
        'content': content, 'document_html': ''.join(blocks), 'tags': 'Word 文档',
        'source': source_name, 'builtin': bool(builtin),
        'sensitive': bool(SENSITIVE_PATTERN.search(content)),
        'created_at': now, 'updated_at': now,
    }


def extract_document_entries(path, password=None, builtin=False):
    extension = os.path.splitext(path)[1].lower()
    if extension == '.xlsx':
        return extract_workbook_entries(path, password, builtin=builtin, source=os.path.basename(path))
    if extension == '.docx':
        return [extract_word_entry(path, builtin=builtin, source=os.path.basename(path))]
    return organize_content(
        extract_document_text(path), source=os.path.basename(path), builtin=builtin,
        file_type=file_type_for_path(path),
    )


def _open_workbook(path, password=None):
    from openpyxl import load_workbook
    try:
        return load_workbook(path, read_only=False, data_only=True)
    except (zipfile.BadZipFile, OSError):
        if not password:
            raise PermissionError('PASSWORD_REQUIRED')
        import msoffcrypto
        decrypted = io.BytesIO()
        with open(path, 'rb') as stream:
            office = msoffcrypto.OfficeFile(stream)
            office.load_key(password=password)
            office.decrypt(decrypted)
        decrypted.seek(0)
        try:
            return load_workbook(decrypted, read_only=False, data_only=True)
        except (zipfile.BadZipFile, OSError, ValueError) as exc:
            raise ValueError('工作簿密码不正确或文件无法读取') from exc


def _cell_text(value):
    if value is None:
        return ''
    if isinstance(value, (datetime.datetime, datetime.date, datetime.time)):
        return value.isoformat(sep=' ') if isinstance(value, datetime.datetime) else value.isoformat()
    return str(value).strip()


def extract_workbook_entries(path, password=None, builtin=False, source=None):
    """Return one structured knowledge entry per non-empty worksheet."""
    from openpyxl.utils import get_column_letter

    workbook = _open_workbook(path, password)
    now = datetime.datetime.now().isoformat(timespec='seconds')
    source_name = source or os.path.basename(path)
    entries = []
    try:
        for sheet in workbook.worksheets:
            # Some supplied workbooks contain formatting down to row 1,048,500.
            # Determine the real rectangle from non-empty cells instead of max_row.
            nonempty = [
                (row, column, cell)
                for (row, column), cell in sheet._cells.items()
                if _cell_text(cell.value)
            ]
            if not nonempty:
                continue
            max_row = max(item[0] for item in nonempty)
            max_column = max(item[1] for item in nonempty)
            rows = [
                [_cell_text(sheet.cell(row=row, column=column).value) for column in range(1, max_column + 1)]
                for row in range(1, max_row + 1)
            ]
            widths = []
            for column in range(1, max_column + 1):
                width = sheet.column_dimensions[get_column_letter(column)].width
                widths.append(round(min(max(float(width or 10), 5), 60), 1))

            header_rows = []
            cell_styles = {}
            for row, column, cell in nonempty:
                fill_rgb = getattr(getattr(cell.fill, 'fgColor', None), 'rgb', None)
                fill_rgb = fill_rgb[-6:] if isinstance(fill_rgb, str) and len(fill_rgb) >= 6 else ''
                bold = bool(cell.font and cell.font.bold)
                if row <= 30 and (bold or (fill_rgb and fill_rgb not in ('000000', 'FFFFFF'))):
                    header_rows.append(row - 1)
                if bold or (fill_rgb and fill_rgb not in ('000000', 'FFFFFF')):
                    style = {}
                    if bold:
                        style['bold'] = True
                    if fill_rgb and fill_rgb not in ('000000', 'FFFFFF'):
                        style['background'] = f'#{fill_rgb}'
                    cell_styles[f'{row - 1},{column - 1}'] = style
            if not header_rows:
                header_rows = [next(index for index, row in enumerate(rows) if any(row))]

            searchable = '\n'.join('\t'.join(row) for row in rows)
            category = classify_content(f'{source_name}\n{sheet.title}\n{searchable}')
            digest_input = json.dumps([source_name, sheet.title, rows], ensure_ascii=False, separators=(',', ':'))
            digest = hashlib.sha256(digest_input.encode('utf-8')).hexdigest()[:20]
            entries.append({
                'id': f'seed-sheet-{digest}' if builtin else uuid.uuid4().hex,
                'title': f'{os.path.splitext(source_name)[0]} · {sheet.title}',
                'category': category,
                'content_type': 'workbook_sheet',
                'file_type': 'EXCEL',
                'content': '',
                'tags': f'Excel 工作表 {sheet.title}',
                'source': source_name,
                'sheet_name': sheet.title,
                'rows': rows,
                'column_widths': widths,
                'header_rows': sorted(set(header_rows)),
                'cell_styles': cell_styles,
                'row_count': max_row,
                'column_count': max_column,
                'builtin': bool(builtin),
                'sensitive': bool(SENSITIVE_PATTERN.search(searchable)),
                'created_at': now,
                'updated_at': now,
            })
    finally:
        workbook.close()
    return entries


def extract_workbook_text(path, password=None):
    """Compatibility helper for callers that explicitly need plain text."""
    entries = extract_workbook_entries(path, password)
    blocks = []
    for entry in entries:
        blocks.append(
            f"Excel: {entry['source']}\nSheet: {entry['sheet_name']}\n"
            + '\n'.join('\t'.join(row) for row in entry['rows'])
        )
    return '\n\n================================================================\n\n'.join(blocks)


def split_content(text):
    normalized = str(text or '').replace('\r\n', '\n').replace('\r', '\n').strip()
    if not normalized:
        return []
    sections = re.split(r'(?m)^\s*(?:={12,}|-{20,}|\*{20,})\s*$', normalized)
    result = []
    for section in sections:
        section = section.strip()
        if not section:
            continue
        if len(section) <= 6500:
            result.append(section)
            continue
        paragraphs = [item.strip() for item in re.split(r'\n{3,}', section) if item.strip()]
        current = []
        current_size = 0
        for paragraph in paragraphs or [section]:
            if current and current_size + len(paragraph) > 5000:
                result.append('\n\n'.join(current))
                current, current_size = [], 0
            current.append(paragraph)
            current_size += len(paragraph)
        if current:
            result.append('\n\n'.join(current))
    return result


def classify_content(content):
    lowered = content.casefold()
    scores = {
        category: sum(lowered.count(keyword.casefold()) for keyword in keywords)
        for category, keywords in CATEGORY_KEYWORDS.items()
    }
    if not any(scores.values()):
        return 'other'
    # Database connection blocks often contain credentials; keep their operational
    # subject as database instead of reducing everything sensitive to access.
    if scores['database'] and ('jdbc:' in lowered or 'service_name' in lowered or 'oracle' in lowered):
        scores['database'] += 3
    return max(scores, key=scores.get)


def suggest_title(content, category, sequence=1):
    lines = [re.sub(r'^\s*[#>*\-]+\s*', '', line).strip() for line in content.splitlines()]
    candidates = [line for line in lines if line]
    for line in candidates[:8]:
        if len(line) <= 58 and not re.search(r'(?i)password|密码|jdbc:|https?://|\bselect\b', line):
            return line
    label = CATEGORIES.get(category, CATEGORIES['other'])
    if candidates:
        snippet = re.sub(r'\s+', ' ', candidates[0])[:32]
        return f'{label} · {snippet}'
    return f'{label} {sequence:03d}'


def organize_content(text, source='直接粘贴', builtin=False, file_type=None):
    now = datetime.datetime.now().isoformat(timespec='seconds')
    entries = []
    for index, content in enumerate(split_content(text), 1):
        category = classify_content(content)
        digest = hashlib.sha256(content.encode('utf-8')).hexdigest()[:20]
        entries.append({
            'id': f'seed-{digest}' if builtin else uuid.uuid4().hex,
            'title': suggest_title(content, category, index),
            'category': category,
            'content_type': 'text_document',
            'file_type': file_type or file_type_for_path(source),
            'content': content,
            'tags': '',
            'source': source,
            'builtin': bool(builtin),
            'sensitive': bool(SENSITIVE_PATTERN.search(content)),
            'created_at': now,
            'updated_at': now,
        })
    return entries


def load_seed_entries():
    """加载内置说明种子。

    发布包仅允许安全空模板；禁止再把真实环境账密打进 resources。
    用户私有笔记请走 data/private_knowledge.json。
    """
    entries = []
    for filename, source in (('private_knowledge_seed.txt', '内置说明（安全模板）'),):
        path = _resource_path(filename)
        if os.path.exists(path):
            try:
                text = read_text_file(path)
            except (OSError, ValueError, TypeError):
                continue
            # 空/纯注释则跳过，避免无意义条目
            body = '\n'.join(
                line for line in text.splitlines()
                if line.strip() and not line.strip().startswith('#')
            ).strip()
            if body:
                entries.extend(organize_content(body, source=source, builtin=True))
    workbook_seed = _resource_path('private_knowledge_seed_workbooks.json')
    try:
        with open(workbook_seed, 'r', encoding='utf-8') as stream:
            loaded = json.load(stream)
        if isinstance(loaded, list):
            for entry in loaded:
                if not isinstance(entry, dict):
                    continue
                # 跳过明显敏感内置表（防止误放回仓库）
                blob = ' '.join(
                    str(entry.get(k, '')) for k in ('title', 'source', 'tags', 'content')
                )
                rows = entry.get('rows') or []
                if rows:
                    blob += ' ' + ' '.join(
                        str(cell) for row in rows[:30] for cell in (row or [])[:12]
                    )
                if SENSITIVE_PATTERN.search(blob) and entry.get('builtin', True):
                    # 含密码等模式的内置表一律不加载
                    continue
                entries.append(dict(entry, builtin=True))
    except (OSError, ValueError, TypeError):
        pass
    return entries


def load_custom_entries(path=None):
    target = path or PRIVATE_KNOWLEDGE_FILE
    try:
        with open(target, 'r', encoding='utf-8') as stream:
            loaded = json.load(stream)
        if not isinstance(loaded, list):
            return []
        return [
            entry for entry in loaded
            if isinstance(entry, dict) and (entry.get('content') or entry.get('rows'))
        ]
    except (OSError, ValueError, TypeError):
        return []


def save_custom_entries(entries, path=None):
    target = path or PRIVATE_KNOWLEDGE_FILE
    if path is None:
        ensure_config_dir()
    else:
        os.makedirs(os.path.dirname(os.path.abspath(target)), exist_ok=True)
    custom = [dict(entry, builtin=False) for entry in entries if not entry.get('builtin')]
    with open(target, 'w', encoding='utf-8') as stream:
        json.dump(custom, stream, ensure_ascii=False, indent=2)


def entry_fingerprint(entry):
    if entry.get('content_type') == 'workbook_sheet':
        value = [entry.get('source', ''), entry.get('sheet_name', ''), entry.get('rows', [])]
        return hashlib.sha256(json.dumps(value, ensure_ascii=False, separators=(',', ':')).encode('utf-8')).hexdigest()
    return hashlib.sha256(str(entry.get('content', '')).strip().encode('utf-8')).hexdigest()


def export_workbook_entry(entry, path, row_indexes=None, column_indexes=None):
    """Export all or a visible subset of one structured worksheet."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter

    rows = entry.get('rows', [])
    row_indexes = list(range(len(rows))) if row_indexes is None else list(row_indexes)
    column_count = entry.get('column_count') or max((len(row) for row in rows), default=0)
    column_indexes = list(range(column_count)) if column_indexes is None else list(column_indexes)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = str(entry.get('sheet_name') or 'Sheet1')[:31]
    styles = entry.get('cell_styles', {})
    for output_row, source_row in enumerate(row_indexes, 1):
        source_values = rows[source_row] if source_row < len(rows) else []
        for output_column, source_column in enumerate(column_indexes, 1):
            value = source_values[source_column] if source_column < len(source_values) else ''
            cell = sheet.cell(output_row, output_column, value=value or None)
            style = styles.get(f'{source_row},{source_column}', {})
            if style.get('bold'):
                cell.font = Font(bold=True)
            background = str(style.get('background', '')).lstrip('#')
            if len(background) == 6:
                cell.fill = PatternFill('solid', fgColor=background)
    widths = entry.get('column_widths', [])
    for output_column, source_column in enumerate(column_indexes, 1):
        if source_column < len(widths):
            sheet.column_dimensions[get_column_letter(output_column)].width = widths[source_column]
    header_rows = set(entry.get('header_rows', []))
    if row_indexes and row_indexes[0] in header_rows:
        sheet.freeze_panes = 'A2'
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    workbook.save(path)
    workbook.close()
    return path


def export_word_entry(entry, path):
    document = Document()
    blocks = []
    document_html = str(entry.get('document_html', ''))
    if document_html:
        class WordHtmlParser(HTMLParser):
            def __init__(self):
                super().__init__(); self.blocks = []; self.paragraph = None; self.table = None; self.row = None; self.cell = None; self.skip = 0
            def handle_starttag(self, tag, _attrs):
                tag = tag.lower()
                if tag in ('style', 'script', 'head'): self.skip += 1
                elif tag == 'table': self.table = []
                elif tag == 'tr' and self.table is not None: self.row = []
                elif tag in ('td', 'th') and self.row is not None: self.cell = []
                elif tag == 'p' and self.cell is None: self.paragraph = []
                elif tag == 'br':
                    target = self.cell if self.cell is not None else self.paragraph
                    if target is not None: target.append('\n')
            def handle_endtag(self, tag):
                tag = tag.lower()
                if tag in ('style', 'script', 'head') and self.skip: self.skip -= 1
                elif tag in ('td', 'th') and self.cell is not None:
                    self.row.append(''.join(self.cell).strip()); self.cell = None
                elif tag == 'tr' and self.row is not None:
                    self.table.append(self.row); self.row = None
                elif tag == 'table' and self.table is not None:
                    self.blocks.append(('table', self.table)); self.table = None
                elif tag == 'p' and self.paragraph is not None:
                    self.blocks.append(('paragraph', ''.join(self.paragraph).strip())); self.paragraph = None
            def handle_data(self, data):
                if self.skip: return
                target = self.cell if self.cell is not None else self.paragraph
                if target is not None: target.append(data)
        parser = WordHtmlParser(); parser.feed(document_html); blocks = parser.blocks
    if blocks:
        for kind, value in blocks:
            if kind == 'paragraph':
                document.add_paragraph(value)
            elif value:
                column_count = max((len(row) for row in value), default=1)
                table = document.add_table(rows=len(value), cols=column_count)
                table.style = 'Table Grid'
                for row_index, row in enumerate(value):
                    for column_index, cell_value in enumerate(row):
                        table.cell(row_index, column_index).text = cell_value
    else:
        for line in str(entry.get('content', '')).splitlines() or ['']:
            document.add_paragraph(line)
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    document.save(path)
    return path


def search_entries(entries, query='', category='all'):
    """支持中文原文 / 全拼 / 首字母；索引不含密钥。置顶条目排在前面。"""
    from tools.list_pin import sort_with_pin
    from tools.pinyin_search import build_search_blob, match_query
    result = []
    for entry in entries:
        if category != 'all' and entry.get('category') != category:
            continue
        parts = [str(entry.get(key, '')) for key in ('title', 'content', 'tags', 'source', 'sheet_name')]
        # 表格行只取有限单元格，避免敏感大报文进索引
        for row in (entry.get('rows', []) or [])[:80]:
            parts.extend(str(value) for value in (row or [])[:16])
        blob = build_search_blob(*parts)
        if match_query(blob, query):
            result.append(entry)
    return sort_with_pin(
        result,
        secondary_key=lambda e: (
            str(e.get('updated_at') or e.get('created_at') or ''),
            str(e.get('title') or ''),
        ),
    )
