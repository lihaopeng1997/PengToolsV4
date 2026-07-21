# -*- coding: utf-8 -*-
"""生成 docs/user_guide/PengToolsHub_使用说明_V4.27.docx（与内置 HTML 说明对齐）。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'user_guide' / 'PengToolsHub_使用说明_V4.27.docx'


def set_run_font(run, size=11, bold=False, color=None, name='微软雅黑'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else (14 if level == 2 else 12),
            bold=True,
            color=(79, 115, 95),
        )
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(it)
        set_run_font(run, size=11)


def add_nums(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(it)
        set_run_font(run, size=11)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(2.4)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(t.add_run('PengToolsHub 使用说明'), size=22, bold=True, color=(79, 115, 95))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub.add_run('V4 Private · 作者 Lihp · Windows 离线桌面工具台'), size=11, color=(92, 101, 96))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        meta.add_run('数据仅存本机 · 无账号 · 无云同步 · 与软件内置 HTML 版内容一致'),
        size=10,
        color=(92, 101, 96),
    )
    doc.add_paragraph()

    add_heading(doc, '1. 产品简介与安装', 1)
    add_p(
        doc,
        'PengToolsHub 是面向内网办公的离线 Windows 工具集合，把日常开发、运维、需求归档、SQL 发版、接口排查、加解密等流程收敛到一个本地窗口。',
    )
    add_heading(doc, '1.1 主要能力', 2)
    add_bullets(doc, [
        '管理需求 / BUG、SVN 工作副本与文件库',
        '整理 SQL、生成发版 Excel',
        '写日报；（彩蛋解锁后）自我学习资料库',
        '网关 SM2+SM4 加解密、JSON/XML 等格式工具',
        '本机 HTTP/HTTPS 抓包、按环境请求测试、明细导出导入',
    ])
    add_heading(doc, '1.2 安装启动', 2)
    add_nums(doc, [
        '解压 PengToolsHub_Private_Offline_Setup.zip',
        '运行 PengToolsHub_Private.exe',
        '首次运行在 EXE 同级创建 data/ 存放配置与业务数据',
    ])
    add_p(doc, '支持单实例：重复启动会激活已打开窗口。')

    add_heading(doc, '2. 界面总览', 1)
    add_heading(doc, '2.1 左侧导航', 2)
    add_bullets(doc, [
        '工作台：首页（最近需求、待升级、快捷入口）',
        '交付管理：需求管理、发版联动、接口文档更新、日报',
        '开发工具：加解密、接口排查、格式工具、证件、VIN、运维助手',
        '个人：自我学习（需彩蛋解锁）',
        '左下角：设置；用户菜单（语言 / 悬浮栏 / 关于 / 使用说明 / 退出）',
    ])
    add_heading(doc, '2.2 悬浮栏', 2)
    add_p(doc, '快捷键 Ctrl+Shift+P 打开或收起桌面悬浮入口，可在设置中编辑快捷模块。')

    add_heading(doc, '3. 首页工作台', 1)
    add_bullets(doc, [
        '最近需求：按更新时间展示，点击跳转并定位',
        '待升级事项：仅「已填上线日期且日期≥今天」、未实际上线、非已上线/关闭/取消/暂停',
        '今天上线：高亮并优先展示',
        '常用工具：一键跳转各模块',
    ])

    add_heading(doc, '4. 需求管理', 1)
    add_heading(doc, '4.1 需求树', 2)
    add_bullets(doc, [
        '按月份/类型组织需求与 BUG',
        '搜索支持拼音，可筛选系统/类型/状态',
        '批量删除有二次确认，默认焦点在取消',
    ])
    add_heading(doc, '4.2 文件库', 2)
    add_bullets(doc, [
        '列：名称、类型、时间、大小、路径；仅名称列带图标',
        '可拖列宽、横向滚动、表头调序',
        '支持拖入本地文件到当前需求工作副本（并尝试 svn add）',
    ])
    add_heading(doc, '4.3 SVN', 2)
    add_bullets(doc, [
        '检出/更新/提交/锁定等依赖本机 SVN 命令行',
        '使用本机缓存认证，软件不保存密码',
        '内网环境请在现场验证',
    ])

    add_heading(doc, '5. 发版联动（升级准备）', 1)
    add_nums(doc, [
        '确认升级日期',
        '勾选参与发版的需求/BUG',
        '多系统分别整理 SQL',
        '生成发版 Excel（内置模板 23 列表头固定）',
    ])
    add_p(doc, '开发分支 SVN 可为空，不阻塞生成。验证 SQL 不进入 SVN 提交目录逻辑。')

    add_heading(doc, '6. 日报', 1)
    add_bullets(doc, [
        '左侧历史日期，右侧编辑完成/风险/计划/备注',
        '未保存切换日期：草稿保留，可能显示「未保存」',
        '「复制为今日」：把当前内容写成今天草稿，再点保存',
        '需求「写入日报」追加模板，不覆盖已写内容',
        '提醒时间在设置中配置',
    ])

    add_heading(doc, '7. 网关加解密', 1)
    add_bullets(doc, [
        'SM2+SM4 网关报文解密与 JSON 查看',
        '从接口排查送入时可自动带报文与 Key',
        '密钥/明文/报文默认不落盘、不写日志',
    ])

    add_heading(doc, '8. 接口排查（Private）', 1)
    add_heading(doc, '8.1 开始抓包', 2)
    add_nums(doc, [
        '本机 127.0.0.1:端口 启动 MITM 代理',
        '临时修改 Windows 系统代理指向该地址',
        '请用浏览器访问业务页（必要时完全重启浏览器）',
    ])
    p = doc.add_paragraph()
    run = p.add_run(
        '重要：抓包会改系统代理。请用「停止抓包」或正常退出以自动恢复。强杀进程可能导致代理残留。'
    )
    set_run_font(run, size=11, bold=True, color=(180, 35, 24))

    add_heading(doc, '8.2 停止与恢复', 2)
    add_bullets(doc, [
        '停止抓包：恢复系统代理，会话列表默认保留',
        '清空：才清空内存抓包记录',
        '「恢复系统代理」：异常后一键修复',
        '软件启动时也会自动清理残留抓包代理',
    ])
    add_heading(doc, '8.3 请求测试', 2)
    add_nums(doc, [
        '维护环境 Base：http(s)://host:port（可多环境保存）',
        '从会话填充：替换 host，保留 path/query；Body 优先解密明文',
        '编辑 Method/Headers/Params/Body 后发送',
    ])
    add_heading(doc, '8.4 导出导入', 2)
    add_bullets(doc, [
        '格式 pengtools_iface_session_v1（JSON）',
        '含 URL 与优先解密的请求/响应',
        '可导入或拖入请求测试页自动填充',
    ])
    add_p(doc, '抓包报文/Cookie/Token 只存内存；配置仅存端口、环境、证书指纹等。')

    add_heading(doc, '9. 格式工具与其它模块', 1)
    add_bullets(doc, [
        '格式工具：JSON/XML/SQL/Base64/URL/时间戳/Java 堆栈等离线处理',
        '接口文档更新：SQL 驱动 DOCX',
        '证件类型 / VIN：测试数据',
        '运维助手：命令查询复制，不自动执行破坏性命令',
        '自我学习：彩蛋解锁，解锁状态存 data/settings.json',
    ])

    add_heading(doc, '10. 设置与关于', 1)
    add_bullets(doc, [
        '主题（含夜间）、字体、语言',
        '悬浮栏透明度、置顶、快捷入口',
        '关闭行为：询问 / 托盘 / 退出',
        '日报提醒',
        '用户菜单：关于、使用说明、退出',
    ])

    add_heading(doc, '11. 数据目录与升级', 1)
    add_bullets(doc, [
        '开发：PengToolsV4/data/',
        '安装运行：EXE 同级 data/',
        '常见文件：settings.json、requirements.json、daily_reports.json、interface_debug.json',
        '升级只替换 EXE 等程序文件，不要删除 data 目录',
    ])

    add_heading(doc, '12. 常见问题 FAQ', 1)
    add_heading(doc, 'Q1 抓包列表为空？', 2)
    add_bullets(doc, [
        '确认已开始抓包',
        '完全退出浏览器后重开再访问业务页',
        'HTTPS 需信任本机抓包证书',
        '检查筛选是否过严',
    ])
    add_heading(doc, 'Q2 抓包后其它接口不好使？', 2)
    add_bullets(doc, [
        '点停止抓包或「恢复系统代理」',
        '重启 PengToolsHub（启动会自动清理）',
        '或在 Windows 代理设置中关闭系统代理',
    ])
    add_heading(doc, 'Q3 请求测试失败？', 2)
    add_bullets(doc, [
        '检查环境 Base 格式与目标服务可达性',
        '查看响应区错误信息',
    ])
    add_heading(doc, 'Q4 日报内容丢失？', 2)
    add_p(doc, '请点「保存日报」。未保存草稿仅进程内保留；退出后只保留已写入 JSON 的内容。')
    add_heading(doc, 'Q5 发版 Excel 列异常？', 2)
    add_p(doc, '请使用软件内置模板生成，勿改动 23 列表头顺序。')

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(end.add_run('— 全文完 · PengToolsHub · Lihp —'), size=10, color=(92, 101, 96))

    doc.save(str(OUT))
    print(f'wrote {OUT} ({OUT.stat().st_size} bytes)')


if __name__ == '__main__':
    main()
